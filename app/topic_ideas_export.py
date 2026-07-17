from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def _clean(value: Any) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip()


def _list(value: Any) -> list[Any]:
    if isinstance(value, list):
        return [item for item in value if item not in (None, "")]
    return [value] if value not in (None, "") else []


def _safe_filename(value: str) -> str:
    text = re.sub(r"[^A-Za-z0-9._-]+", "_", _clean(value)).strip("._")
    return (text[:80] or "topic_ideas") + ".docx"


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _add_hyperlink(paragraph, text: str, url: str) -> None:
    clean_url = _clean(url)
    if not clean_url:
        return
    relationship_id = paragraph.part.relate_to(
        clean_url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    colour = OxmlElement("w:color")
    colour.set(qn("w:val"), "1D4ED8")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.append(colour)
    run_properties.append(underline)
    run.append(run_properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _add_labelled_paragraph(doc: Document, label: str, value: Any) -> None:
    text = _clean(value)
    if not text:
        return
    paragraph = doc.add_paragraph()
    label_run = paragraph.add_run(f"{label}: ")
    label_run.bold = True
    paragraph.add_run(text)


def _add_bullets(doc: Document, values: Any, numbered: bool = False) -> None:
    style = "List Number" if numbered else "List Bullet"
    for value in _list(values):
        text = _clean(value)
        if text:
            doc.add_paragraph(text, style=style)


def _add_resource(doc: Document, source: dict[str, Any], *, instrument: bool = False) -> None:
    title = _clean(source.get("title") if instrument else source.get("name"))
    if not title:
        return
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(title)
    run.bold = True
    meta_parts: list[str] = []
    if instrument:
        authors = source.get("authors") or []
        if isinstance(authors, list):
            authors = ", ".join(_clean(item) for item in authors if _clean(item))
        if _clean(authors):
            meta_parts.append(_clean(authors))
        if _clean(source.get("year")):
            meta_parts.append(f"({_clean(source.get('year'))})")
        if _clean(source.get("source")):
            meta_parts.append(_clean(source.get("source")))
    else:
        if _clean(source.get("provider") or source.get("discovery_database")):
            meta_parts.append(_clean(source.get("provider") or source.get("discovery_database")))
        if _clean(source.get("year")):
            meta_parts.append(_clean(source.get("year")))
    if meta_parts:
        doc.add_paragraph(" · ".join(meta_parts))
    description = _clean(source.get("candidate_use") if instrument else source.get("description"))
    if description:
        doc.add_paragraph(description)
    matched = source.get("matched_constructs") if instrument else source.get("matched_variables_or_constructs")
    matched_text = ", ".join(_clean(item) for item in _list(matched) if _clean(item))
    if matched_text:
        _add_labelled_paragraph(doc, "Matched constructs", matched_text)
    note = _clean(source.get("access_and_adaptation_note") if instrument else source.get("access_note"))
    if note:
        p = doc.add_paragraph()
        r = p.add_run(note)
        r.italic = True
        r.font.color.rgb = RGBColor(180, 83, 9)
    url = _clean(source.get("url") or (f"https://doi.org/{_clean(source.get('doi'))}" if source.get("doi") else ""))
    if url:
        p = doc.add_paragraph()
        _add_hyperlink(p, "Open source", url)


def export_topic_ideas_docx(result: dict[str, Any], out_dir: Path) -> Path:
    out_dir.mkdir(parents=True, exist_ok=True)
    research_area = _clean(result.get("research_area") or result.get("query") or "Topic Ideas")
    path = out_dir / _safe_filename(f"{research_area}_topic_ideas")

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        styles[style_name].font.name = "Aptos Display"
        styles[style_name].font.color.rgb = RGBColor(15, 23, 42)

    title = doc.add_paragraph()
    title.style = styles["Title"]
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("ProjectReady AI Topic Ideas")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(research_area)
    subtitle_run.bold = True
    subtitle_run.font.size = Pt(13)

    doc.add_paragraph(
        "Editable research-topic planning output. Verify the feasibility, sources, instruments, data access, ethics and institutional requirements with a supervisor before adopting any idea."
    )

    meta = doc.add_table(rows=0, cols=2)
    meta.style = "Table Grid"
    metadata = [
        ("Academic level", result.get("selected_level")),
        ("Access", "Free preview" if result.get("free_preview") else "Unlocked set"),
        ("Ideas returned", result.get("ideas_returned") or len(result.get("ideas") or [])),
        ("Recent-reference window", result.get("recent_reference_window")),
        ("Search query", result.get("query")),
    ]
    for label, value in metadata:
        if _clean(value):
            cells = meta.add_row().cells
            cells[0].text = label
            cells[1].text = _clean(value)
            _set_cell_shading(cells[0], "EAF2FF")
            for run in cells[0].paragraphs[0].runs:
                run.bold = True

    if _clean(result.get("trend_summary")):
        doc.add_heading("Trend and gap summary", level=1)
        doc.add_paragraph(_clean(result.get("trend_summary")))

    ideas = _list(result.get("ideas"))[:12]
    for index, raw_idea in enumerate(ideas, start=1):
        if not isinstance(raw_idea, dict):
            continue
        idea = raw_idea
        doc.add_page_break() if index > 1 else None
        doc.add_heading(f"{index}. {_clean(idea.get('title')) or 'Untitled idea'}", level=1)
        if _clean(idea.get("synopsis")):
            doc.add_paragraph(_clean(idea.get("synopsis")))

        objectives = idea.get("proposed_objectives") or {}
        if isinstance(objectives, dict):
            doc.add_heading("Proposed research objectives", level=2)
            _add_labelled_paragraph(doc, "General objective", objectives.get("general_objective"))
            specific = _list(objectives.get("specific_objectives"))
            if specific:
                p = doc.add_paragraph()
                p.add_run("Specific objectives").bold = True
                _add_bullets(doc, specific, numbered=True)
            _add_labelled_paragraph(doc, "Level alignment", objectives.get("level_alignment"))

        _add_labelled_paragraph(doc, "Current trend or gap", idea.get("current_research_trend_or_gap"))
        _add_labelled_paragraph(doc, "Possible methodology", idea.get("possible_methodology"))

        variables = _list(idea.get("possible_variables_or_constructs"))
        if variables:
            doc.add_heading("Variables or constructs", level=2)
            _add_bullets(doc, variables)

        data_directions = _list(idea.get("possible_data_sources"))
        if data_directions:
            doc.add_heading("Topic-specific data direction", level=2)
            _add_bullets(doc, data_directions)

        guidance = idea.get("research_resource_guidance") or {}
        if isinstance(guidance, dict):
            secondary = _list(guidance.get("secondary_data_sources"))
            instruments = _list(guidance.get("questionnaire_or_instrument_sources"))
            if secondary:
                doc.add_heading("Strongly matched secondary-data candidates", level=2)
                for source in secondary:
                    if isinstance(source, dict):
                        _add_resource(doc, source, instrument=False)
            if instruments:
                doc.add_heading("Strongly matched instrument candidates", level=2)
                for source in instruments:
                    if isinstance(source, dict):
                        _add_resource(doc, source, instrument=True)

        _add_labelled_paragraph(doc, "Potential contribution", idea.get("potential_contribution"))
        attention = _clean(idea.get("attention_note"))
        if attention:
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(f"Attention: {attention}")
            run.bold = True
            run.font.color.rgb = RGBColor(180, 83, 9)

    sources = _list(result.get("source_records_used"))
    if sources:
        doc.add_page_break()
        doc.add_heading("Literature records used for trend grounding", level=1)
        source_number = 0
        for source in sources:
            if not isinstance(source, dict):
                continue
            source_number += 1
            p = doc.add_paragraph()
            title_text = _clean(source.get("title")) or "Untitled source"
            p.add_run(f"{source_number}. {title_text}").bold = True
            authors = source.get("authors") or []
            if isinstance(authors, list):
                authors = ", ".join(_clean(item) for item in authors if _clean(item))
            details = " · ".join(
                item for item in [
                    _clean(authors),
                    _clean(source.get("year")),
                    _clean(source.get("source") or source.get("database")),
                ] if item
            )
            if details:
                doc.add_paragraph(details)
            if _clean(source.get("url")):
                link_p = doc.add_paragraph()
                _add_hyperlink(link_p, "Open literature record", _clean(source.get("url")))

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("ProjectReady AI · Topic ideas require independent academic review and verification.")

    doc.save(path)
    return path
