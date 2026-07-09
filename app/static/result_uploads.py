from __future__ import annotations

import csv
import io
import re
from pathlib import Path
from typing import Any

from docx import Document

MAX_EXTRACTED_CHARS = 60000
MAX_PREVIEW_CHARS = 1800
MAX_ROWS_PER_SHEET = 80
MAX_COLUMNS_PER_SHEET = 20


def extract_result_file(filename: str, content: bytes) -> dict[str, Any]:
    """Extract readable text from a results/analysis output file.

    Supported formats are intentionally common for student projects:
    txt/md/csv/tsv, docx, xlsx/xlsm, and pdf. Unsupported files are decoded
    as text where possible.
    """
    clean_name = Path(filename or "results_upload").name
    suffix = Path(clean_name).suffix.lower()

    if not content:
        raise ValueError("The uploaded file is empty.")

    if suffix in {".txt", ".md", ".log"}:
        text = _decode_text(content)
    elif suffix in {".csv", ".tsv"}:
        text = _extract_delimited(content, delimiter="\t" if suffix == ".tsv" else ",")
    elif suffix == ".docx":
        text = _extract_docx(content)
    elif suffix in {".xlsx", ".xlsm"}:
        text = _extract_xlsx(content)
    elif suffix == ".pdf":
        text = _extract_pdf(content)
    else:
        text = _decode_text(content)

    text = _clean_text(text)
    if not text.strip():
        raise ValueError("No readable text could be extracted from the uploaded file.")

    truncated = len(text) > MAX_EXTRACTED_CHARS
    extracted_text = text[:MAX_EXTRACTED_CHARS]
    return {
        "filename": clean_name,
        "file_type": suffix.lstrip(".") or "unknown",
        "characters_extracted": len(extracted_text),
        "truncated": truncated,
        "extracted_text": extracted_text,
        "preview": extracted_text[:MAX_PREVIEW_CHARS],
    }


def _decode_text(content: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="replace")


def _extract_delimited(content: bytes, delimiter: str = ",") -> str:
    text = _decode_text(content)
    reader = csv.reader(io.StringIO(text), delimiter=delimiter)
    rows: list[list[str]] = []
    for idx, row in enumerate(reader):
        if idx >= MAX_ROWS_PER_SHEET:
            break
        rows.append([str(cell).strip() for cell in row[:MAX_COLUMNS_PER_SHEET]])

    if not rows:
        return text

    max_cols = max(len(row) for row in rows)
    normalised = [row + [""] * (max_cols - len(row)) for row in rows]
    return _rows_to_markdown_table(normalised, title="Tabular analysis evidence")


def _extract_docx(content: bytes) -> str:
    doc = Document(io.BytesIO(content))
    parts: list[str] = []

    for para in doc.paragraphs:
        value = para.text.strip()
        if value:
            parts.append(value)

    for table_index, table in enumerate(doc.tables, 1):
        rows: list[list[str]] = []
        for row in table.rows[:MAX_ROWS_PER_SHEET]:
            rows.append([cell.text.strip().replace("\n", " ") for cell in row.cells[:MAX_COLUMNS_PER_SHEET]])
        if rows:
            parts.append(_rows_to_markdown_table(rows, title=f"DOCX table {table_index}"))

    return "\n\n".join(parts)


def _extract_xlsx(content: bytes) -> str:
    try:
        from openpyxl import load_workbook
    except Exception as exc:
        raise ValueError("Excel uploads require openpyxl. Add openpyxl to requirements.txt and redeploy.") from exc

    wb = load_workbook(io.BytesIO(content), data_only=True, read_only=True)
    parts: list[str] = []

    for ws in wb.worksheets:
        rows: list[list[str]] = []
        for row_idx, row in enumerate(ws.iter_rows(values_only=True), 1):
            if row_idx > MAX_ROWS_PER_SHEET:
                break
            values = ["" if cell is None else str(cell) for cell in row[:MAX_COLUMNS_PER_SHEET]]
            if any(value.strip() for value in values):
                rows.append(values)
        if rows:
            parts.append(_rows_to_markdown_table(rows, title=f"Excel sheet: {ws.title}"))

    return "\n\n".join(parts)


def _extract_pdf(content: bytes) -> str:
    try:
        from pypdf import PdfReader
    except Exception as exc:
        raise ValueError("PDF uploads require pypdf. Add pypdf to requirements.txt and redeploy.") from exc

    reader = PdfReader(io.BytesIO(content))
    parts: list[str] = []
    for page_index, page in enumerate(reader.pages[:30], 1):
        page_text = page.extract_text() or ""
        if page_text.strip():
            parts.append(f"[PDF page {page_index}]\n{page_text.strip()}")
    return "\n\n".join(parts)


def _rows_to_markdown_table(rows: list[list[str]], title: str = "Table") -> str:
    if not rows:
        return ""
    max_cols = max(len(row) for row in rows)
    rows = [row + [""] * (max_cols - len(row)) for row in rows]

    header = rows[0]
    # If the first row looks numeric or blank, create generic headers.
    if not any(str(cell).strip() for cell in header) or sum(_looks_numeric(cell) for cell in header) > max(1, len(header) // 2):
        header = [f"Column {i + 1}" for i in range(max_cols)]
        body = rows
    else:
        body = rows[1:]

    lines = [f"{title}", "", "| " + " | ".join(_escape_cell(cell) for cell in header) + " |"]
    lines.append("| " + " | ".join("---" for _ in header) + " |")
    for row in body:
        lines.append("| " + " | ".join(_escape_cell(cell) for cell in row) + " |")
    return "\n".join(lines)


def _escape_cell(value: Any) -> str:
    text = str(value or "").strip().replace("|", "/")
    return re.sub(r"\s+", " ", text)


def _looks_numeric(value: Any) -> bool:
    text = str(value or "").strip().replace(",", "")
    if not text:
        return False
    try:
        float(text)
        return True
    except ValueError:
        return False


def _clean_text(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()
