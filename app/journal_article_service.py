from __future__ import annotations

import io
import json
import os
import re
from datetime import datetime
from typing import Any

try:
    from app.source_finder import search_literature_sources
except Exception:  # pragma: no cover
    search_literature_sources = None

MAX_SOURCE_CONTEXT = 18

_RETRACTION_TERMS = re.compile(
    r"\b(retracted|retraction\s+notice|withdrawn|removed\s+article|expression\s+of\s+concern|erratum\s+to\s+retracted)\b",
    flags=re.IGNORECASE,
)

_ATTENTION_RE = re.compile(
    r"\[(?:insert|verify|confirm|provide|supply|complete|replace|check|add|update|obtain|state|specify|include)\b[^\]]*\]",
    flags=re.IGNORECASE,
)


def _safe_get_openai_client():
    api_key = os.getenv("OPENAI_API_KEY", "").strip()
    if not api_key:
        return None
    try:
        from openai import OpenAI
        return OpenAI(api_key=api_key)
    except Exception:
        return None


def _select_article_model(level: str, article_type: str = "") -> str:
    """Route journal article drafting by academic depth and publication complexity."""
    level_l = (level or "").strip().lower()
    type_l = (article_type or "").strip().lower()
    is_doctoral = any(token in level_l for token in ["phd", "doctor", "dba", "ded", "professional doctorate"])
    is_research_masters = "research masters" in level_l or "mphil" in level_l
    is_review_article = any(token in type_l for token in ["systematic", "scoping", "meta-analysis", "review article", "conceptual"])
    if is_doctoral:
        return os.getenv("OPENAI_ARTICLE_DOCTORAL_MODEL", os.getenv("OPENAI_DOCTORAL_DRAFT_MODEL", "gpt-5.5")).strip()
    if is_research_masters or is_review_article:
        return os.getenv("OPENAI_ARTICLE_RESEARCH_MODEL", os.getenv("OPENAI_RESEARCH_MASTERS_DRAFT_MODEL", "gpt-5.5")).strip()
    if "non-research" in level_l or "master" in level_l:
        return os.getenv("OPENAI_ARTICLE_MASTERS_MODEL", "gpt-5.4").strip()
    return os.getenv("OPENAI_ARTICLE_BACHELOR_MODEL", os.getenv("OPENAI_BACHELOR_DRAFT_MODEL", "gpt-5.4")).strip()


def _looks_retracted(src: dict[str, Any]) -> bool:
    fields = [
        src.get("title"), src.get("type"), src.get("subtype"), src.get("status"), src.get("publication_status"),
        src.get("update_type"), src.get("relation_type"), src.get("abstract"), src.get("note"), src.get("warning"),
    ]
    combined = " ".join(str(x or "") for x in fields)
    if _RETRACTION_TERMS.search(combined):
        return True
    flags = ["is_retracted", "retracted", "has_retraction", "is_withdrawn", "withdrawn", "removed", "expression_of_concern"]
    return any(bool(src.get(flag)) for flag in flags)


def _build_search_profile(payload: dict[str, Any]) -> dict[str, Any]:
    objectives = []
    for raw in str(payload.get("objectives") or "").split("\n"):
        item = raw.strip(" -;,")
        if item:
            objectives.append(item)
    return {
        "title": str(payload.get("article_title") or "").strip(),
        "research_area": str(payload.get("research_area") or payload.get("article_title") or "").strip(),
        "study_context": str(payload.get("context") or "").strip(),
        "level": str(payload.get("academic_level") or "Research Masters (e.g. MPhil)").strip(),
        "research_approach": str(payload.get("methodology") or "").strip(),
        "data_type": str(payload.get("article_type") or "").strip(),
        "objectives": objectives[:8],
        "notes": " ".join([
            str(payload.get("target_journal") or ""),
            str(payload.get("variables_constructs") or ""),
            str(payload.get("key_findings") or ""),
        ]).strip(),
    }


def _search_sources(payload: dict[str, Any]) -> tuple[list[dict[str, Any]], list[dict[str, Any]], dict[str, Any]]:
    if not payload.get("include_source_search", True) or search_literature_sources is None:
        return [], [], {"provider_errors": [], "query": ""}
    profile = _build_search_profile(payload)
    query = " ".join([
        str(payload.get("article_title") or ""),
        str(payload.get("research_area") or ""),
        str(payload.get("context") or ""),
        str(payload.get("variables_constructs") or ""),
        str(payload.get("theory_or_framework") or ""),
        str(payload.get("key_findings") or ""),
    ]).strip()
    result = search_literature_sources(
        profile=profile,
        query=query,
        max_results=int(os.getenv("PROJECTREADY_ARTICLE_SOURCE_LIMIT", "24")),
        include_older_foundational=bool(payload.get("include_older_foundational", True)),
    )
    raw_sources = result.get("sources") or []
    blocked = [s for s in raw_sources if _looks_retracted(s)]
    usable = [s for s in raw_sources if not _looks_retracted(s)]
    return usable, blocked, result


def _source_context(sources: list[dict[str, Any]]) -> list[dict[str, Any]]:
    records = []
    for idx, src in enumerate(sources[:MAX_SOURCE_CONTEXT], start=1):
        authors = src.get("authors") or []
        if isinstance(authors, str):
            authors = [authors]
        records.append({
            "key": f"S{idx}",
            "title": src.get("title", ""),
            "authors": authors,
            "year": src.get("year", ""),
            "source": src.get("source", ""),
            "doi": src.get("doi", ""),
            "url": src.get("url", ""),
            "abstract": str(src.get("abstract") or "")[:1200],
            "database": src.get("database", ""),
            "relevance_tier": src.get("relevance_tier", ""),
            "citation_count": src.get("citation_count", ""),
            "reference_entry_hint": src.get("apa_hint") or src.get("reference_entry_hint") or "",
        })
    return records


def _extract_text(response: Any) -> str:
    return str(getattr(response, "output_text", "") or "").strip()


def _strip_code_fences(text: str) -> str:
    if not text:
        return ""
    return re.sub(r"^```(?:markdown|md)?\s*|\s*```$", "", text.strip(), flags=re.IGNORECASE | re.MULTILINE).strip()


def _fallback_article(payload: dict[str, Any], sources: list[dict[str, Any]]) -> str:
    title = str(payload.get("article_title") or "Article Draft").strip()
    article_type = str(payload.get("article_type") or "Empirical research article").strip()
    citation_style = str(payload.get("citation_style") or "APA 7th").strip()
    source_note = ""
    if sources:
        source_note = "\n\nThe draft should be strengthened with the following reviewed source records: " + "; ".join(
            f"S{i+1}: {s.get('title', 'Untitled')} ({s.get('year', 'n.d.')})" for i, s in enumerate(sources[:6])
        )
    return f"""# {title}

## Article Type and Target

Article type: {article_type}. Target journal: {payload.get('target_journal') or '[insert target journal]'}.

## Abstract

[insert 180-250 word structured or unstructured abstract after confirming the final results, sample, method and contribution]

## Keywords

[insert 4-6 keywords]

## 1. Introduction

This section should establish the research problem, the current scholarly conversation, the context of the study and the article's contribution. Add recent evidence and verified citations before submission.{source_note}

## 2. Literature Review and Theoretical Positioning

[insert focused literature synthesis using verified current sources and foundational theory where appropriate]

## 3. Methods

[insert article-ready methods section, including design, setting, population/sample, data source, measures, analysis technique, validity/reliability or trustworthiness, and ethics]

## 4. Results

[insert analysed results, tables and figures. Do not invent statistics, coefficients, themes or p-values]

## 5. Discussion

[insert interpretation of findings against theory, prior studies and the study context]

## 6. Conclusion

[insert concise conclusion, contribution, limitations and future research]

## Declarations

Funding: [confirm funding statement]

Conflict of interest: [confirm conflict-of-interest statement]

Ethics approval: [confirm ethics approval or exemption]

Data availability: [confirm data availability statement]

## References

[insert {citation_style} references for sources cited in the article only]
""".strip()


def _finalise_article_text(text: str) -> str:
    text = _strip_code_fences(text or "")
    text = re.sub(r"<span\s+[^>]*>(.*?)</span>", r"\1", text, flags=re.I | re.S)
    text = text.replace("—", ", ").replace(" – ", ", ").replace("–", "-").replace("‑", "-")
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    return text.strip()


def draft_journal_article(payload: dict[str, Any]) -> dict[str, Any]:
    if not str(payload.get("article_title") or "").strip():
        raise ValueError("Article title or working topic is required.")
    sources, blocked, search_result = _search_sources(payload)
    source_records = _source_context(sources)
    model = _select_article_model(str(payload.get("academic_level") or ""), str(payload.get("article_type") or ""))
    client = _safe_get_openai_client()
    provider_errors = search_result.get("provider_errors") or [] if isinstance(search_result, dict) else []

    if not client or os.getenv("PROJECTREADY_ARTICLE_USE_AI", "1").strip().lower() in {"0", "false", "no"}:
        article_text = _fallback_article(payload, sources)
        mode = "metadata_fallback"
    else:
        current_year = datetime.now().year
        prompt = {
            "task": "Draft a publishable journal article manuscript from the user's inputs and current scholarly metadata.",
            "article_inputs": payload,
            "current_year": current_year,
            "source_records": source_records,
            "strict_rules": [
                "Use the target journal guidelines supplied by the user as formatting and structural rules. If they are missing, use a standard scholarly article structure appropriate to the article type.",
                "Do not guarantee publication. Produce a journal-article-ready draft that still requires author verification, supervisor review and journal formatting checks.",
                "Do not fabricate results, sample sizes, p-values, coefficients, themes, quotations, ethics approvals, funding, conflicts of interest, datasets or citations.",
                "Use bracketed attention placeholders where details are missing, for example [confirm sample size], [insert regression table], [verify ethics approval] or [insert target journal word limit].",
                "Use only supplied source records, user-provided reference notes or sources that can be stated with confidence. Do not invent references.",
                "Never use retracted, withdrawn, removed or expression-of-concern sources to support any argument, table, citation or reference entry.",
                "Keep citations accurate and include a References section containing only cited sources.",
                "Respect the selected citation style and target journal notes where possible.",
                "Write in polished formal British English with clear, publishable argument, but avoid exaggerated claims and promotional language.",
                "Minimise em dashes and en dashes; use commas, semicolons, colons, parentheses or separate sentences instead.",
            ],
            "recommended_article_structures": {
                "empirical": ["Title", "Abstract", "Keywords", "Introduction", "Literature Review/Theory", "Methods", "Results", "Discussion", "Conclusion", "Declarations", "References"],
                "systematic_review": ["Title", "Abstract", "Keywords", "Introduction", "Methods", "Results", "Discussion", "Conclusion", "References"],
                "conceptual": ["Title", "Abstract", "Keywords", "Introduction", "Conceptual/Theoretical Background", "Proposed Framework", "Discussion", "Implications", "Conclusion", "References"],
                "case_study": ["Title", "Abstract", "Keywords", "Introduction", "Case Context", "Methods", "Findings", "Discussion", "Conclusion", "References"],
            },
            "output_format": [
                "Return Markdown only.",
                "Use clean numbered headings where appropriate.",
                "Use tables only when they improve clarity and keep them compact.",
                "Include an Article Readiness Checklist at the end with missing items and actions.",
            ],
        }
        try:
            response = client.responses.create(
                model=model,
                instructions=(
                    "You are ProjectReady AI's journal article drafting assistant. Draft publishable-quality academic manuscripts from verified inputs. "
                    "Follow journal guidelines when supplied. Use current, non-retracted source metadata for literature framing. "
                    "Do not invent evidence, citations, results or declarations. Use bracketed attention placeholders for missing information."
                ),
                input=json.dumps(prompt, ensure_ascii=False, indent=2),
            )
            article_text = _extract_text(response) or _fallback_article(payload, sources)
            mode = "ai_draft"
        except Exception as exc:
            provider_errors = provider_errors + [f"OpenAI article drafting failed: {str(exc)[:180]}"]
            article_text = _fallback_article(payload, sources)
            mode = "metadata_fallback_after_ai_error"

    article_text = _finalise_article_text(article_text)
    return {
        "article_text": article_text,
        "model_used": model if client else "none",
        "mode": mode,
        "source_records_used": source_records,
        "excluded_retracted_count": len(blocked),
        "excluded_retracted_titles": [str(s.get("title") or "Untitled") for s in blocked[:10]],
        "provider_errors": provider_errors,
        "quality_filters": [
            "Retracted, withdrawn, removed and expression-of-concern records excluded where detectable in metadata.",
            "References limited to cited and verified/supplied sources.",
            "Missing article details rendered as bracketed attention placeholders.",
        ],
    }


def _add_inline_runs(paragraph, text: str) -> None:
    """Add basic bold/italic and attention placeholder styling to a paragraph."""
    from docx.shared import RGBColor
    pos = 0
    token_re = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|\[[^\]]+\])")
    for match in token_re.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos:match.start()])
        token = match.group(0)
        run_text = token
        bold = False
        italic = False
        if token.startswith("**") and token.endswith("**"):
            run_text = token[2:-2]
            bold = True
        elif token.startswith("*") and token.endswith("*"):
            run_text = token[1:-1]
            italic = True
        run = paragraph.add_run(run_text)
        run.bold = bold
        run.italic = italic
        if _ATTENTION_RE.fullmatch(token):
            run.font.color.rgb = RGBColor(192, 0, 0)
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _add_markdown_table(doc, lines: list[str]) -> None:
    rows = []
    for line in lines:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        if cells and not all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in cells):
            rows.append(cells)
    if not rows:
        return
    width = max(len(r) for r in rows)
    table = doc.add_table(rows=0, cols=width)
    table.style = "Table Grid"
    for row_idx, cells in enumerate(rows):
        row = table.add_row().cells
        for i in range(width):
            row[i].text = cells[i] if i < len(cells) else ""
        if row_idx == 0:
            for cell in row:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True


def export_article_docx(article_text: str, title: str = "Journal Article Draft") -> tuple[io.BytesIO, str]:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    safe_title = re.sub(r"[^A-Za-z0-9_-]+", "_", (title or "journal_article")[:80]).strip("_") or "journal_article"
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5

    table_buffer: list[str] = []
    for raw_line in _finalise_article_text(article_text).splitlines():
        line = raw_line.rstrip()
        if line.strip().startswith("|") and line.strip().endswith("|"):
            table_buffer.append(line)
            continue
        if table_buffer:
            _add_markdown_table(doc, table_buffer)
            table_buffer = []
        if not line.strip():
            continue
        if line.startswith("# "):
            p = doc.add_heading(line[2:].strip(), level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif re.match(r"^[-*•]\s+", line):
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_runs(p, re.sub(r"^[-*•]\s+", "", line).strip())
        elif re.match(r"^\d+[.)]\s+", line):
            p = doc.add_paragraph(style="List Number")
            _add_inline_runs(p, re.sub(r"^\d+[.)]\s+", "", line).strip())
        else:
            p = doc.add_paragraph()
            _add_inline_runs(p, line)
    if table_buffer:
        _add_markdown_table(doc, table_buffer)
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream, f"{safe_title}_journal_article_draft.docx"
