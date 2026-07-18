from __future__ import annotations

import difflib
import io
import json
import os
import re
from datetime import datetime
from typing import Any

from app.source_finder import search_literature_sources
from app.action_items import detach_action_items
from app.ai_service import (
    _clean_chapter_references,
    _ensure_markdown_heading_spacing,
    _normalise_objectives_and_questions,
    _normalise_purpose_of_study,
)
from app.scholarly_humanizer import (
    analyse_scholarly_style,
    build_humanizer_batches,
    humanize_scholarly_text,
    humanizer_variation_profile,
    scholarly_humanizer_prompt_rules,
    validate_humanizer_preservation,
    variation_targets_met,
)

_REVISION_BLUE = (0, 112, 192)
_ACTION_RED = (192, 0, 0)

_RETRACTION_TERMS = re.compile(
    r"\b(retracted|retraction\s+notice|withdrawn|removed\s+article|expression\s+of\s+concern|erratum\s+to\s+retracted)\b",
    flags=re.IGNORECASE,
)

_ATTENTION_RE = re.compile(
    r"\[(?:insert|verify|confirm|provide|supply|complete|replace|check|add|update|obtain|state|specify|include|revise|review|conduct|perform|run|collect|clarify|report|resolve|address|identify|upload|attach|calculate|test|assess|determine|seek|action|required|author\s+action|student\s+action|supervisor\s+action)\b[^\]]*\]",
    flags=re.IGNORECASE,
)

_ACTION_LABEL_RE = re.compile(
    r"\b(?:action required|author action|student action|supervisor action|required action|remaining action|attention required|user action)\s*:\s*",
    flags=re.IGNORECASE,
)

CHAPTER_PAGE_TARGETS: dict[str, dict[str, tuple[int, int]]] = {
    "bachelors": {
        "introduction": (10, 15),
        "literature_review": (15, 22),
        "methodology": (10, 15),
        "results_discussion": (20, 25),
        "conclusion": (8, 12),
        "other": (8, 15),
    },
    "non_research_masters": {
        "introduction": (10, 15),
        "literature_review": (20, 30),
        "methodology": (12, 18),
        "results_discussion": (20, 30),
        "conclusion": (8, 15),
        "other": (10, 18),
    },
    "research_masters": {
        "introduction": (15, 20),
        "literature_review": (35, 45),
        "methodology": (15, 22),
        "results_discussion": (20, 32),
        "conclusion": (8, 12),
        "other": (12, 22),
    },
    "professional_doctorate": {
        "introduction": (15, 22),
        "literature_review": (40, 60),
        "methodology": (25, 35),
        "results_discussion": (35, 45),
        "conclusion": (10, 15),
        "other": (15, 28),
    },
    "phd": {
        "introduction": (25, 35),
        "literature_review": (60, 80),
        "methodology": (30, 45),
        "results_discussion": (60, 80),
        "conclusion": (20, 30),
        "other": (20, 35),
    },
}

CITATION_DENSITY_TARGETS: dict[str, dict[str, tuple[int, int]]] = {
    "bachelors": {
        "introduction": (12, 16), "literature_review": (16, 22),
        "methodology": (6, 9), "results_discussion": (6, 10),
        "conclusion": (4, 7), "other": (7, 10),
    },
    "non_research_masters": {
        "introduction": (13, 18), "literature_review": (18, 24),
        "methodology": (7, 10), "results_discussion": (7, 11),
        "conclusion": (5, 8), "other": (8, 11),
    },
    "research_masters": {
        "introduction": (15, 20), "literature_review": (20, 28),
        "methodology": (8, 12), "results_discussion": (8, 13),
        "conclusion": (6, 9), "other": (9, 13),
    },
    "professional_doctorate": {
        "introduction": (16, 22), "literature_review": (22, 30),
        "methodology": (9, 13), "results_discussion": (9, 14),
        "conclusion": (7, 10), "other": (10, 14),
    },
    "phd": {
        "introduction": (18, 24), "literature_review": (24, 32),
        "methodology": (10, 15), "results_discussion": (10, 16),
        "conclusion": (8, 12), "other": (11, 16),
    },
}


def _env_bool(name: str, default: bool = False) -> bool:
    value = os.getenv(name)
    if value is None:
        return default
    return value.strip().lower() not in {"0", "false", "no", "off", ""}


def _safe_get_openai_client():
    api_key = os.getenv("OPENAI_API_KEY", "").strip()
    if not api_key:
        return None
    try:
        from openai import OpenAI
        return OpenAI(
            api_key=api_key,
            timeout=float(os.getenv("PROJECTREADY_CHAPTER_REVISION_TIMEOUT_SECONDS", "900")),
        )
    except Exception:
        return None




def _env_int(name: str, default: int, minimum: int | None = None, maximum: int | None = None) -> int:
    try:
        value = int(os.getenv(name, str(default)) or default)
    except Exception:
        value = default
    if minimum is not None:
        value = max(minimum, value)
    if maximum is not None:
        value = min(maximum, value)
    return value


def _format_provider_errors(errors: list[Any]) -> str:
    rows: list[str] = []
    for item in errors:
        if isinstance(item, dict):
            provider = str(item.get("provider") or "provider").strip()
            error = str(item.get("error") or item.get("message") or "unavailable").strip()
            if "HTTP Error 429" in error or "rate limit" in error.lower():
                error = "temporary rate limit from this scholarly metadata provider"
            rows.append(f"- {provider}: {error[:220]}")
        else:
            rows.append(f"- {str(item)[:240]}")
    return "\n".join(rows) or "- The AI revision service was unavailable."


def _revision_model_candidates(level: str) -> list[str]:
    primary = _revision_model(level)
    candidates = [
        primary,
        os.getenv("OPENAI_CHAPTER_REVISION_FALLBACK_MODEL", ""),
        os.getenv("OPENAI_FINAL_SYNTHESIS_MODEL", ""),
        os.getenv("OPENAI_SECTION_ANALYSIS_MODEL", ""),
        os.getenv("OPENAI_MODEL", ""),
    ]
    # Keep conservative fallback names last. Unsupported models fail fast and the
    # next configured option is tried. Production should still set the explicit
    # env vars above for the models available to the account.
    if _level_key(level) in {"research_masters", "professional_doctorate", "phd"}:
        candidates.extend(["gpt-5.6-sol", "gpt-5.6-terra", "gpt-5.6-luna"])
    else:
        candidates.extend(["gpt-5.6-terra", "gpt-5.6-luna"])
    result: list[str] = []
    seen: set[str] = set()
    for value in candidates:
        item = str(value or "").strip()
        if not item or item in seen:
            continue
        seen.add(item)
        result.append(item)
    return result


def _response_output_text(response: Any) -> str:
    output = str(getattr(response, "output_text", "") or "").strip()
    if output:
        return output
    try:
        chunks: list[str] = []
        for item in getattr(response, "output", []) or []:
            for content in getattr(item, "content", []) or []:
                text = getattr(content, "text", None)
                if text:
                    chunks.append(str(text))
        return "\n".join(chunks).strip()
    except Exception:
        return ""


def _call_responses_with_fallbacks(
    client: Any,
    *,
    level: str,
    prompt: dict[str, Any],
    instructions: str,
    max_output_tokens: int,
) -> tuple[str, str, list[dict[str, str]]]:
    """Call the revision model with retries and configured fallbacks.

    The old implementation made one very large request. A single timeout then
    returned the chapter unchanged and consumed the user's revision entitlement.
    This helper tries configured fallback models and reduces the token ceiling on
    retry so a transient timeout does not become a completed fallback report.
    """
    errors: list[dict[str, str]] = []
    attempts_per_model = _env_int("PROJECTREADY_CHAPTER_REVISION_MODEL_ATTEMPTS", 2, minimum=1, maximum=4)
    for candidate in _revision_model_candidates(level):
        token_budget = max_output_tokens
        for attempt in range(1, attempts_per_model + 1):
            try:
                response = client.responses.create(
                    model=candidate,
                    max_output_tokens=max(3500, token_budget),
                    instructions=instructions,
                    input=json.dumps(prompt, ensure_ascii=False, indent=2),
                )
                output = _response_output_text(response)
                if output:
                    return output, candidate, errors
                errors.append({
                    "provider": "openai",
                    "error": f"{candidate} returned an empty response on attempt {attempt}.",
                })
            except Exception as exc:
                message = str(exc)[:240]
                errors.append({
                    "provider": "openai",
                    "error": f"{candidate} attempt {attempt}: {message}",
                })
                if re.search(r"timed?\s*out|timeout|request\s+timed\s+out", message, flags=re.IGNORECASE):
                    token_budget = max(3500, int(token_budget * 0.55))
                    continue
                break
    raise RuntimeError("Chapter revision failed after configured model retries. " + _format_provider_errors(errors))


def _level_key(level: str) -> str:
    value = (level or "").strip().lower()
    if "phd" in value:
        return "phd"
    if any(token in value for token in ["professional doctorate", "dba", "ded", "doctorate"]):
        return "professional_doctorate"
    if "non-research" in value or "non research" in value:
        return "non_research_masters"
    if "research masters" in value or "mphil" in value:
        return "research_masters"
    return "bachelors"


def _chapter_key(chapter_type: str) -> str:
    value = (chapter_type or "").strip().lower()
    if "literature" in value or "chapter two" in value or "chapter 2" in value:
        return "literature_review"
    if "method" in value or "chapter three" in value or "chapter 3" in value:
        return "methodology"
    if "result" in value or "discussion" in value or "chapter four" in value or "chapter 4" in value:
        return "results_discussion"
    if any(token in value for token in ["summary", "conclusion", "recommendation", "chapter five", "chapter 5"]):
        return "conclusion"
    if "introduction" in value or "chapter one" in value or "chapter 1" in value:
        return "introduction"
    return "other"


def _page_target(level: str, chapter_type: str) -> tuple[int, int]:
    return CHAPTER_PAGE_TARGETS[_level_key(level)][_chapter_key(chapter_type)]


_SECTION_PAGE_TARGETS: dict[str, tuple[int, int]] = {
    "bachelors": (2, 4),
    "non_research_masters": (3, 5),
    "research_masters": (4, 7),
    "professional_doctorate": (5, 9),
    "phd": (6, 10),
}


def _resolved_page_target(payload: dict[str, Any], level: str, chapter_type: str) -> tuple[int, int]:
    default_min, default_max = _page_target(level, chapter_type)
    if bool(payload.get("custom_target_pages_enabled")):
        try:
            custom_min = max(1, min(int(payload.get("target_page_min") or default_min), 120))
            custom_max = max(1, min(int(payload.get("target_page_max") or default_max), 120))
        except Exception:
            custom_min, custom_max = default_min, default_max
        if custom_max < custom_min:
            custom_min, custom_max = custom_max, custom_min
        return custom_min, custom_max

    if str(payload.get("strengthening_scope") or "whole_chapter") == "selected_sections":
        section_count = len(payload.get("selected_section_titles") or [])
        section_count += len(payload.get("new_section_titles") or [])
        section_count += len(payload.get("custom_new_sections") or [])
        section_count = max(1, section_count)
        per_min, per_max = _SECTION_PAGE_TARGETS[_level_key(level)]
        return max(1, min(default_max, section_count * per_min)), max(2, min(default_max, section_count * per_max))
    return default_min, default_max


_CHAPTER_WORD_NUMBERS = {
    "one": 1, "two": 2, "three": 3, "four": 4, "five": 5,
    "six": 6, "seven": 7, "eight": 8, "nine": 9, "ten": 10,
}


def _chapter_heading_number(line: str) -> int | None:
    value = re.sub(r"^\s*#{1,6}\s*", "", str(line or "")).strip()
    value = re.sub(r"\s+", " ", value)
    match = re.match(r"^chapter\s+(one|two|three|four|five|six|seven|eight|nine|ten|\d{1,2})\b", value, flags=re.IGNORECASE)
    if match:
        token = match.group(1).lower()
        return int(token) if token.isdigit() else _CHAPTER_WORD_NUMBERS.get(token)
    match = re.match(r"^(\d{1,2})(?:\.0)?\s+(?!\d)([A-Za-z][^\n]{2,100})$", value)
    if match and "." not in match.group(1):
        return int(match.group(1))
    return None


def _scope_to_selected_chapter(text: str, chapter_type: str, uploaded_content_scope: str) -> tuple[str, dict[str, Any]]:
    value = _finalise_chapter_text(text)
    if uploaded_content_scope != "complete_thesis":
        return value, {"uploaded_content_scope": "selected_chapter", "chapter_isolated": False, "reason": "selected chapter supplied"}

    chapter_number_match = re.match(r"^\s*(\d+)", str(chapter_type or ""))
    chapter_number = int(chapter_number_match.group(1)) if chapter_number_match else None
    if not chapter_number:
        raise ValueError(
            "A complete thesis can only be isolated when a numbered chapter is selected. "
            "Choose Chapters One to Five or paste only the custom chapter text."
        )

    lines = value.splitlines(keepends=True)
    offsets: list[int] = []
    cursor = 0
    headings: list[tuple[int, int]] = []
    for line in lines:
        offsets.append(cursor)
        number = _chapter_heading_number(line)
        if number is not None:
            headings.append((cursor, number))
        cursor += len(line)

    candidates: list[tuple[int, str]] = []
    for index, (start, number) in enumerate(headings):
        if number != chapter_number:
            continue
        end = len(value)
        for next_start, next_number in headings[index + 1:]:
            if next_number != chapter_number:
                end = next_start
                break
        segment = value[start:end].strip()
        if len(segment) >= 100:
            candidates.append((len(segment), segment))

    if not candidates:
        raise ValueError(
            f"The uploaded complete thesis could not be safely separated at Chapter {chapter_number}. "
            "Make sure the document contains a clear heading such as 'CHAPTER TWO' or '2 LITERATURE REVIEW', "
            "or upload/paste only the selected chapter."
        )
    _, selected = max(candidates, key=lambda item: item[0])
    return selected, {
        "uploaded_content_scope": "complete_thesis",
        "chapter_isolated": True,
        "selected_chapter_number": chapter_number,
        "original_character_count": len(value),
        "selected_character_count": len(selected),
    }


def _citation_target(level: str, chapter_type: str) -> tuple[int, int]:
    return CITATION_DENSITY_TARGETS[_level_key(level)][_chapter_key(chapter_type)]


def _revision_model(level: str) -> str:
    key = _level_key(level)
    specific = {
        "bachelors": "OPENAI_CHAPTER_REVISION_BACHELOR_MODEL",
        "non_research_masters": "OPENAI_CHAPTER_REVISION_MASTERS_MODEL",
        "research_masters": "OPENAI_CHAPTER_REVISION_RESEARCH_MODEL",
        "professional_doctorate": "OPENAI_CHAPTER_REVISION_DOCTORAL_MODEL",
        "phd": "OPENAI_CHAPTER_REVISION_DOCTORAL_MODEL",
    }[key]
    return (
        os.getenv(specific)
        or os.getenv("OPENAI_CHAPTER_REVISION_MODEL")
        or ("gpt-5.6-sol" if key in {"professional_doctorate", "phd"} else "gpt-5.6-terra")
    ).strip()


def _looks_retracted(source: dict[str, Any]) -> bool:
    fields = [
        source.get("title"), source.get("type"), source.get("status"),
        source.get("publication_status"), source.get("abstract"),
        source.get("retraction_status"),
    ]
    if _RETRACTION_TERMS.search(" ".join(str(item or "") for item in fields)):
        return True
    return any(bool(source.get(flag)) for flag in [
        "is_retracted", "retracted", "has_retraction", "is_withdrawn",
        "withdrawn", "removed", "expression_of_concern",
    ])


def _source_identity(source: dict[str, Any]) -> str:
    doi = str(source.get("doi") or "").strip().lower()
    if doi:
        return f"doi:{doi}"
    title = re.sub(r"[^a-z0-9]+", " ", str(source.get("title") or "").lower()).strip()
    return f"title:{title}"


def _source_context(sources: list[dict[str, Any]]) -> list[dict[str, Any]]:
    result = []
    for index, source in enumerate(sources[:60], start=1):
        result.append({
            "key": f"S{index}",
            "title": source.get("title", ""),
            "authors": source.get("authors", []),
            "year": source.get("year", ""),
            "source": source.get("source", ""),
            "doi": source.get("doi", ""),
            "url": source.get("url", ""),
            "abstract": str(source.get("abstract") or "")[:900],
            "database": source.get("database", ""),
            "apa_hint": source.get("apa_hint", ""),
        })
    return result


def _search_sources(payload: dict[str, Any]) -> tuple[list[dict[str, Any]], list[dict[str, Any]], dict[str, Any]]:
    attached = [item for item in (payload.get("source_bank") or []) if isinstance(item, dict)]
    blocked = [item for item in attached if _looks_retracted(item)]
    safe = [item for item in attached if not _looks_retracted(item)]
    search_result: dict[str, Any] = {"provider_errors": [], "sources": []}

    if bool(payload.get("include_source_search", True)):
        profile = {
            "title": str(payload.get("thesis_title") or payload.get("chapter_title") or "").strip(),
            "research_area": str(payload.get("research_area") or "").strip(),
            "study_context": str(payload.get("context") or "").strip(),
            "discipline": str(payload.get("discipline") or "").strip(),
            "data_type": str(payload.get("chapter_type") or "").strip(),
            "objectives": [
                line.strip(" -;")
                for line in re.split(r"[\n;]+", str(payload.get("objectives") or ""))
                if line.strip()
            ][:5],
        }
        query = " ".join([
            str(payload.get("source_search_terms") or ""),
            str(payload.get("research_area") or ""),
            str(payload.get("variables_constructs") or ""),
            str(payload.get("theory_framework") or ""),
            str(payload.get("context") or ""),
        ]).strip()
        try:
            search_result = search_literature_sources(
                profile=profile,
                query=query,
                max_results=max(12, min(int(payload.get("source_limit") or 45), 60)),
                include_older_foundational=bool(payload.get("include_older_foundational", True)),
            )
            for source in search_result.get("sources") or []:
                if _looks_retracted(source):
                    blocked.append(source)
                else:
                    safe.append(source)
        except Exception as exc:
            search_result = {
                "provider_errors": [{"provider": "source_search", "error": str(exc)[:220]}],
                "sources": [],
            }

    deduped: list[dict[str, Any]] = []
    seen: set[str] = set()
    for source in safe:
        identity = _source_identity(source)
        if identity in seen:
            continue
        seen.add(identity)
        deduped.append(source)
    return deduped[:60], blocked, search_result


def _finalise_chapter_text(text: str) -> str:
    value = str(text or "").replace("\r\n", "\n").replace("\r", "\n")
    value = re.sub(r"<span[^>]*>", "", value, flags=re.IGNORECASE)
    value = re.sub(r"</span>", "", value, flags=re.IGNORECASE)
    value = re.sub(r"^```(?:markdown|md|text)?\s*|\s*```$", "", value.strip(), flags=re.IGNORECASE | re.MULTILINE)
    value = re.sub(r"[ \t]+\n", "\n", value)
    value = re.sub(r"\n{3,}", "\n\n", value)
    return value.strip()


# Alias retained by the DOCX diff renderer adapted from the ArticleReady module.
_finalise_article_text = _finalise_chapter_text


def _normalise_inline_markdown(text: str) -> str:
    value = str(text or "")
    value = re.sub(r"(?<!\*)\*\*(?!\*)([^*\n]+)\*{3}(?=$|[\s.,;:!?])", r"**\1**", value)
    value = re.sub(r"(?<!\*)\*(?!\*)([^*\n]+)\*{2}(?=$|[\s.,;:!?])", r"*\1*", value)
    value = re.sub(r"(?<!_)__(?!_)([^_\n]+)_{3}(?=$|[\s.,;:!?])", r"__\1__", value)
    value = re.sub(r"(?<!_)_(?!_)([^_\n]+)_{2}(?=$|[\s.,;:!?])", r"_\1_", value)
    return value


def _strip_stray_inline_markers(text: str) -> str:
    value = str(text or "")
    value = re.sub(r"(?<!\w)\*{1,3}(?=\w)", "", value)
    value = re.sub(r"(?<=\w)\*{1,3}(?=$|[\s.,;:!?])", "", value)
    value = re.sub(r"(?<!\w)_{1,3}(?=\w)", "", value)
    value = re.sub(r"(?<=\w)_{1,3}(?=$|[\s.,;:!?])", "", value)
    return value


def _split_action_segments(text: str, bold: bool = False, italic: bool = False) -> list[tuple[str, bool, bool, bool]]:
    value = str(text or "")
    if not value:
        return []
    segments: list[tuple[str, bool, bool, bool]] = []
    position = 0
    for match in re.finditer(r"\[[^\]\n]+\]", value):
        if match.start() > position:
            prefix = value[position:match.start()]
            label = _ACTION_LABEL_RE.search(prefix)
            if label:
                if label.start() > 0:
                    segments.append((prefix[:label.start()], bold, italic, False))
                segments.append((prefix[label.start():], bold, italic, True))
            else:
                segments.append((prefix, bold, italic, False))
        token = match.group(0)
        segments.append((token, bold, italic, bool(_ATTENTION_RE.fullmatch(token))))
        position = match.end()
    if position < len(value):
        tail = value[position:]
        label = _ACTION_LABEL_RE.search(tail)
        if label:
            if label.start() > 0:
                segments.append((tail[:label.start()], bold, italic, False))
            segments.append((tail[label.start():], bold, italic, True))
        else:
            segments.append((tail, bold, italic, False))
    if not segments:
        segments.append((value, bold, italic, bool(_ACTION_LABEL_RE.search(value))))
    return [(part, b, i, action) for part, b, i, action in segments if part]


def _parse_inline_segments(text: str) -> list[tuple[str, bool, bool, bool]]:
    value = _normalise_inline_markdown(text)
    style_re = re.compile(
        r"(?P<bolditalic>\*\*\*(?P<bolditalic_text>.+?)\*\*\*)"
        r"|(?P<bold>\*\*(?P<bold_text>.+?)\*\*)"
        r"|(?P<italic>(?<!\*)\*(?!\*)(?P<italic_text>.+?)(?<!\*)\*(?!\*))"
        r"|(?P<bolditalic_u>___(?P<bolditalic_u_text>.+?)___)"
        r"|(?P<bold_u>__(?P<bold_u_text>.+?)__)"
        r"|(?P<italic_u>(?<!\w)_(?!_)(?P<italic_u_text>.+?)(?<!_)_(?!\w))",
        flags=re.DOTALL,
    )
    segments: list[tuple[str, bool, bool, bool]] = []
    position = 0
    for match in style_re.finditer(value):
        if match.start() > position:
            segments.extend(_split_action_segments(_strip_stray_inline_markers(value[position:match.start()])))
        if match.group("bolditalic") is not None:
            inner, bold, italic = match.group("bolditalic_text"), True, True
        elif match.group("bold") is not None:
            inner, bold, italic = match.group("bold_text"), True, False
        elif match.group("italic") is not None:
            inner, bold, italic = match.group("italic_text"), False, True
        elif match.group("bolditalic_u") is not None:
            inner, bold, italic = match.group("bolditalic_u_text"), True, True
        elif match.group("bold_u") is not None:
            inner, bold, italic = match.group("bold_u_text"), True, False
        else:
            inner, bold, italic = match.group("italic_u_text"), False, True
        segments.extend(_split_action_segments(inner or "", bold=bold, italic=italic))
        position = match.end()
    if position < len(value):
        segments.extend(_split_action_segments(_strip_stray_inline_markers(value[position:])))
    return segments or _split_action_segments(_strip_stray_inline_markers(value))


def _plain_inline_text(text: str) -> str:
    return "".join(segment[0] for segment in _parse_inline_segments(text))


def _strip_code_fences(text: str) -> str:
    return re.sub(r"^```(?:markdown|md)?\s*|\s*```$", "", str(text or "").strip(), flags=re.IGNORECASE | re.MULTILINE).strip()


def _split_revision_package(text: str) -> tuple[str, str, str]:
    raw = _strip_code_fences(text)
    revised_marker = "===REVISED_CHAPTER==="
    report_marker = "===STRENGTHENING_REPORT==="
    matrix_marker = "===SUPERVISOR_RESPONSE_MATRIX==="
    if revised_marker in raw:
        raw = raw.split(revised_marker, 1)[1]
    matrix = ""
    if matrix_marker in raw:
        raw, matrix = raw.split(matrix_marker, 1)
    report = ""
    if report_marker in raw:
        revised, report = raw.split(report_marker, 1)
    else:
        revised = raw
    return revised.strip(), report.strip(), matrix.strip()


def _chapter_rules(chapter_type: str, study_stage: str) -> list[str]:
    key = _chapter_key(chapter_type)
    stage = (study_stage or "").strip().lower()
    common = [
        "Retain the selected chapter identity and its normal thesis or dissertation function.",
        "Preserve valid headings, tables, equations, citations, references and institution-specific terminology.",
        "Do not change the approved topic, objectives, questions, hypotheses, constructs, sample, methods or findings without explicit supporting information.",
    ]
    specific = {
        "introduction": [
            "Strengthen the progression from broad context to the specific problem, gap, purpose, objectives, questions or hypotheses and significance.",
            "Distinguish the practical problem, empirical problem, theoretical or conceptual gap, contextual gap and methodological gap where supported.",
            "Ensure objectives, questions and hypotheses are aligned and do not introduce relationships absent from the study.",
        ],
        "literature_review": [
            "Replace author-by-author listing with synthesis organised around concepts, theories, objectives, themes and relationships.",
            "Clarify definitions, theoretical mechanisms, competing explanations, empirical consistencies, contradictions and unresolved gaps.",
            "Increase citation density with directly relevant sources, but do not pad the chapter or invent bibliographic details.",
        ],
        "methodology": [
            "Check alignment among philosophy, approach, design, population, sampling, measurement, data collection, analysis, validity or trustworthiness and ethics.",
            "Provide methodological justification rather than merely naming procedures.",
            "Keep equations, variable definitions, coding rules and diagnostic requirements consistent with the stated objectives and data.",
            (
                "Use future tense for planned procedures because the study is at proposal stage."
                if "proposal" in stage or "planned" in stage
                else "Use past tense for procedures that were completed, while established methodological principles may remain in present tense."
            ),
        ],
        "results_discussion": [
            "Preserve every confirmed result, coefficient, p-value, confidence interval, theme, quotation, table value and diagnostic outcome.",
            "Do not fabricate missing analysis. Insert a concise red-action placeholder where an essential result or diagnostic is absent.",
            "Separate presentation of results from interpretation where the institutional format requires it, then deepen the discussion by objective.",
            "Compare findings with relevant literature, explain plausible mechanisms, recognise conflicting evidence and avoid causal language unsupported by the design.",
        ],
        "conclusion": [
            "Trace every summary statement, conclusion and recommendation to an actual objective and confirmed finding.",
            "Avoid introducing new results, literature-dependent claims or recommendations that are not supported by the study.",
            "Name responsible actors, implementation conditions, contribution, limitations and future research needs where required.",
        ],
        "other": [
            "Follow the supplied custom chapter purpose, school format and supervisor direction.",
            "Preserve cross-chapter consistency and avoid duplicating content that belongs elsewhere.",
        ],
    }
    return common + specific[key]


def _revision_focus(payload: dict[str, Any]) -> list[str]:
    focus_map = [
        ("strengthen_structure", "chapter structure, section completeness and logical progression"),
        ("strengthen_problem_gap", "problem statement, research gap, purpose and objective alignment"),
        ("strengthen_conceptualisation", "conceptual definitions, theoretical framing, mechanisms and boundary conditions"),
        ("increase_citation_density", "citation density, evidence integration, synthesis and reference relevance"),
        ("assess_method_fit", "alignment of design, sampling, measurement, data collection, analysis, validity or trustworthiness and ethics"),
        ("assess_results", "accuracy, completeness and objective-level reporting of results or qualitative findings"),
        ("deepen_discussion", "interpretive depth, comparison with literature, mechanisms, conflicting evidence, limitations and contribution"),
        ("strengthen_conclusions", "finding-linked conclusions, implications, recommendations, limitations and contribution"),
        ("improve_language", "formal British English, coherence, paragraph development, transitions and natural scholarly rhythm"),
    ]
    return [description for key, description in focus_map if bool(payload.get(key, True))]


def _fallback_report(payload: dict[str, Any], sources: list[dict[str, Any]], errors: list[Any]) -> str:
    source_note = (
        f"{len(sources)} scholarly record(s) were available for relevance screening, but no source was inserted automatically in fallback mode."
        if sources else "No scholarly records were available for source-supported strengthening."
    )
    error_text = _format_provider_errors(errors)
    return f"""# Chapter Strengthening Report

## Status

The existing chapter has been returned without substantive rewriting because the revision model was unavailable. Do not treat this fallback as a completed chapter strengthening.

## Processing notes

{error_text}

## Priority checks

1. Confirm that the uploaded text is the selected chapter.
2. Check all required sections against the school or department format.
3. Align the problem, gap, purpose, objectives, questions, hypotheses, theory, methods, results and conclusions across the thesis.
4. Verify every in-text citation and reference entry.
5. Supply actual data, analysis output, ethics information and institutional details where required.
6. Resolve every bracketed author-action item before submission.

## Source review

{source_note}
""".strip()


def _fallback_matrix(comments: str) -> str:
    items = [line.strip(" -\t") for line in str(comments or "").splitlines() if line.strip()]
    if not items:
        return ""
    rows = ["| Supervisor comment | Revision made | Location | Remaining action |", "|---|---|---|---|"]
    for item in items[:40]:
        rows.append(f"| {item.replace('|', '/')} | [substantive revision not completed in fallback mode] | [identify section] | [action required] |")
    return "\n".join(rows)


def _word_count(text: str) -> int:
    return len(re.findall(r"\b[\w’'-]+\b", _plain_inline_text(text), flags=re.UNICODE))


def _citation_count(text: str) -> int:
    value = str(text or "")
    author_date = re.findall(r"\((?:[^()]*?\b(?:19|20)\d{2}[a-z]?[^()]*)\)", value)
    narrative = re.findall(r"\b[A-Z][A-Za-z'’-]+(?:\s+(?:and|&)\s+[A-Z][A-Za-z'’-]+)?\s*\((?:19|20)\d{2}[a-z]?\)", value)
    numeric = re.findall(r"(?<!\w)\[(?:\d+(?:\s*[-,]\s*\d+)*)\]", value)
    return len(author_date) + len(narrative) + len(numeric)


def _metrics(text: str) -> dict[str, Any]:
    words = _word_count(text)
    citations = _citation_count(text)
    return {
        "word_count": words,
        "estimated_pages": round(words / 350, 1) if words else 0.0,
        "citation_occurrences": citations,
        "citations_per_1000_words": round(citations * 1000 / words, 1) if words else 0.0,
    }




def _humanizer_batch_output_tokens(word_count: int) -> int:
    words = max(250, int(word_count or 0))
    return max(1800, min(9000, int(words * 2.1)))


def _humanize_strengthened_chapter_with_model(
    client: Any,
    model: str,
    text: str,
    *,
    mode: str,
    level: str,
    chapter_type: str,
    payload: dict[str, Any],
) -> str:
    """Apply a protected section-batched model pass after substantive strengthening.

    A failure in this optional style stage never invalidates the completed revision.
    Balanced mode touches only the weakest batches and Deep mode covers all eligible
    batches up to a configurable cap.
    """
    if mode not in {"balanced", "deep"} or not client or not text.strip():
        return text

    threshold = _env_int("PROJECTREADY_HUMANIZER_MODEL_THRESHOLD", 97, minimum=60, maximum=99)
    overall = analyse_scholarly_style(text)
    variation_profile = humanizer_variation_profile()
    style_context = any(str(payload.get(key) or "").strip() for key in (
        "revision_goals", "supervisor_comments", "school_guidelines"
    ))
    if (
        mode == "balanced"
        and not style_context
        and int(overall.get("score") or 100) >= threshold
        and variation_targets_met(overall, variation_profile)
    ):
        return text

    batch_words = _env_int("PROJECTREADY_HUMANIZER_BATCH_WORDS", 1800, minimum=800, maximum=5000)
    batches = build_humanizer_batches(text, max_words=batch_words)
    eligible = [
        index for index, batch in enumerate(batches)
        if not batch.get("protected")
        and int((batch.get("diagnostic") or {}).get("word_count") or 0) >= 120
        and (
            mode == "deep"
            or style_context
            or int((batch.get("diagnostic") or {}).get("score") or 100) < threshold
            or not variation_targets_met(batch.get("diagnostic") or {}, variation_profile)
        )
    ]
    if mode == "balanced":
        eligible.sort(key=lambda index: int((batches[index].get("diagnostic") or {}).get("score") or 100))
        eligible = eligible[:_env_int("PROJECTREADY_HUMANIZER_MAX_BATCHES_BALANCED", 6, minimum=1, maximum=12)]
    else:
        eligible = eligible[:_env_int("PROJECTREADY_HUMANIZER_MAX_BATCHES_DEEP", 16, minimum=1, maximum=24)]
    if not eligible:
        return text

    chosen = set(eligible)
    output: list[str] = []
    for index, batch in enumerate(batches):
        original = str(batch.get("text") or "")
        if index not in chosen:
            output.append(original)
            continue
        prompt = {
            "task": "Refine this strengthened thesis section for natural scholarly flow without changing its substance.",
            "academic_level": level,
            "chapter_type": chapter_type,
            "style_diagnostic": batch.get("diagnostic") or {},
            "variation_profile": variation_profile,
            "revision_direction": str(payload.get("revision_goals") or "").strip(),
            "supervisor_direction": str(payload.get("supervisor_comments") or "").strip(),
            "rules": [
                "Revise rather than restart.",
                *scholarly_humanizer_prompt_rules(),
                "Preserve every heading, citation, reference, date, statistic, objective, research question, hypothesis, table, equation and bracketed action item exactly.",
                "Do not add evidence, citations, findings, examples, interpretations, recommendations or new sections.",
                "Preserve the order of ideas and the strength of claims.",
                "Increase controlled perplexity through context-sensitive lexical and syntactic variety without rare-synonym substitution.",
                "Increase controlled burstiness through a purposeful mix of short, medium and longer synthesis sentences and varied paragraph movement.",
                "Improve directness, sentence movement, paragraph rhythm and logical transitions without rare-synonym substitution.",
                "Keep the word count within six percent of the supplied section.",
                "Return only the revised section with its headings and no report.",
            ],
            "section": original,
        }
        try:
            response = client.responses.create(
                model=model,
                max_output_tokens=_humanizer_batch_output_tokens(int(batch.get("word_count") or 0)),
                instructions="Perform an evidence-preserving, high-variation scholarly naturalness edit. Return only the revised section.",
                input=json.dumps(prompt, ensure_ascii=False, indent=2),
            )
            candidate = _response_output_text(response)
            candidate, _ = humanize_scholarly_text(candidate, mode="balanced") if candidate else (original, {})
            valid, _issues = validate_humanizer_preservation(
                original,
                candidate,
                max_word_change_ratio=float(variation_profile["model_word_change_limit"]),
            )
            output.append(candidate if candidate and valid else original)
        except Exception:
            output.append(original)

    candidate = "\n\n".join(part.strip() for part in output if part.strip()).strip()
    valid, _issues = validate_humanizer_preservation(
        text,
        candidate,
        max_word_change_ratio=float(variation_profile["model_word_change_limit"]),
    )
    return candidate if valid else text


def _revision_long_chapter_strategy(level: str, chapter_type: str, page_min: int, page_max: int) -> dict[str, Any]:
    level_key = _level_key(level)
    chapter_key = _chapter_key(chapter_type)
    target_words = int(round(((page_min + page_max) / 2) * 330))
    enabled = target_words >= int(os.getenv("PROJECTREADY_LONG_CHAPTER_THRESHOLD_WORDS", "12000") or 12000) or (level_key in {"phd", "professional_doctorate"} and chapter_key == "literature_review")
    if not enabled:
        return {"enabled": False, "mode": "standard_strengthening"}
    workflow = [
        "diagnose the current chapter map against the expected doctoral chapter architecture",
        "strengthen the chapter in section batches rather than compressing the whole chapter into a short overview",
        "for literature reviews, check conceptual, theoretical, empirical, methodological, contextual, contradiction, gap and framework coverage separately",
        "insert missing but necessary sections only when justified, and mark them with bracketed confirmation placeholders",
        "merge the strengthened sections for coherence while preserving confirmed student content and citations",
    ]
    return {
        "enabled": True,
        "mode": "long_chapter_staged_strengthening",
        "target_words_estimate": target_words,
        "target_pages": f"{page_min}-{page_max}",
        "workflow": workflow,
        "rules": [
            "Do not rewrite a very long chapter as one compressed summary.",
            "Treat underdeveloped sections as separate strengthening units.",
            "Use bracketed red-action placeholders for missing evidence, sources or supervisor confirmation.",
            "Preserve confirmed content and only deepen it through evidence, synthesis, comparison and alignment checks.",
        ],
    }


def _previous_chapters_revision_context(payload: dict[str, Any]) -> dict[str, Any]:
    """Compact alignment context supplied to the Chapter Strengthener."""
    raw = payload.get("previous_chapters_context") or payload.get("previous_chapters_for_alignment") or ""
    items: list[dict[str, Any]] = []
    if isinstance(raw, dict):
        for item in raw.get("items") or []:
            if not isinstance(item, dict):
                continue
            text = str(item.get("text") or "").strip()
            if text:
                items.append({
                    "label": str(item.get("label") or "Previous chapter context"),
                    "source": str(item.get("source") or "previous_chapter"),
                    "text": text[:16000],
                })
    elif isinstance(raw, list):
        for index, item in enumerate(raw, start=1):
            if isinstance(item, dict):
                text = str(item.get("text") or item.get("extracted_text") or "").strip()
                label = str(item.get("label") or item.get("filename") or f"Previous chapter context {index}")
            else:
                text = str(item or "").strip()
                label = f"Previous chapter context {index}"
            if text:
                items.append({"label": label, "source": "previous_chapter", "text": text[:16000]})
    else:
        text = str(raw or "").strip()
        if text:
            items.append({"label": "Uploaded or pasted previous chapters / full work", "source": "previous_chapter", "text": text[:28000]})

    compact: list[dict[str, Any]] = []
    total = 0
    for item in items:
        text = str(item.get("text") or "").strip()
        if len(text) < 80:
            continue
        remaining = max(0, 52000 - total)
        if remaining <= 0:
            break
        clipped = text[:remaining]
        total += len(clipped)
        compact.append({**item, "text": clipped, "characters": len(clipped)})

    return {
        "available": bool(compact),
        "items": compact,
        "rules": [
            "Use this context to check cross-chapter alignment, not to copy earlier chapters into the revised chapter.",
            "Check consistency of the thesis title, problem, gap, objectives, questions, hypotheses, theory, concepts, variables, methodology, scope and terminology.",
            "Where the current chapter conflicts with earlier chapters, mark the issue as a bracketed attention placeholder instead of silently changing approved study logic.",
            "For Chapter Two and later, ensure the revised chapter connects clearly to the earlier approved study direction.",
        ],
    }


def revise_chapter(payload: dict[str, Any]) -> dict[str, Any]:
    payload = dict(payload)
    raw_uploaded_text = _finalise_chapter_text(str(payload.get("chapter_text") or ""))
    if len(raw_uploaded_text) < 100:
        raise ValueError("Paste or upload the existing thesis or dissertation chapter before requesting strengthening.")

    level = str(payload.get("academic_level") or "Bachelors")
    chapter_type = str(payload.get("chapter_type") or "1. Introduction")
    chapter_text, scope_metadata = _scope_to_selected_chapter(
        raw_uploaded_text,
        chapter_type,
        str(payload.get("uploaded_content_scope") or "selected_chapter"),
    )
    if scope_metadata.get("chapter_isolated"):
        existing_alignment = str(payload.get("previous_chapters_context") or "").strip()
        if len(raw_uploaded_text) <= 52000:
            full_thesis_alignment = raw_uploaded_text
        else:
            middle_start = max(0, (len(raw_uploaded_text) // 2) - 8000)
            full_thesis_alignment = "\n\n".join([
                "[Beginning of complete thesis]\n" + raw_uploaded_text[:18000],
                "[Middle sample of complete thesis]\n" + raw_uploaded_text[middle_start:middle_start + 16000],
                "[End of complete thesis]\n" + raw_uploaded_text[-18000:],
            ])
        payload["previous_chapters_context"] = "\n\n---\n\n".join(
            item for item in [existing_alignment, "Complete thesis supplied for alignment only:\n" + full_thesis_alignment] if item
        )

    strengthening_scope = str(payload.get("strengthening_scope") or "whole_chapter")
    selected_titles = [str(item).strip() for item in payload.get("selected_section_titles") or [] if str(item).strip()]
    new_titles = [str(item).strip() for item in payload.get("new_section_titles") or [] if str(item).strip()]
    custom_new_sections = [item for item in payload.get("custom_new_sections") or [] if isinstance(item, dict) and str(item.get("title") or "").strip()]
    if strengthening_scope == "selected_sections" and not (selected_titles or new_titles or custom_new_sections):
        raise ValueError("Select at least one section to strengthen or add before using selected-sections mode.")

    page_min, page_max = _resolved_page_target(payload, level, chapter_type)
    citation_min, citation_max = _citation_target(level, chapter_type)

    sources, blocked, search_result = _search_sources(payload)
    source_records = _source_context(sources)
    provider_errors = list(search_result.get("provider_errors") or [])
    model = _revision_model(level)
    model_used = model
    client = _safe_get_openai_client()
    ai_enabled = _env_bool("PROJECTREADY_CHAPTER_REVISION_USE_AI", True)

    if not client or not ai_enabled:
        revised_chapter = chapter_text
        strengthening_report = _fallback_report(payload, source_records, provider_errors)
        supervisor_matrix = _fallback_matrix(str(payload.get("supervisor_comments") or ""))
        mode = "metadata_fallback"
    else:
        prompt = {
            "task": (
                "Strengthen an existing thesis or dissertation chapter while preserving confirmed evidence, "
                "the approved study direction and the author's disciplinary voice."
            ),
            "current_year": datetime.now().year,
            "project_profile": {
                "thesis_title": str(payload.get("thesis_title") or "").strip(),
                "chapter_title": str(payload.get("chapter_title") or chapter_type).strip(),
                "chapter_type": chapter_type,
                "academic_level": level,
                "programme_or_discipline": str(payload.get("discipline") or "").strip(),
                "study_stage": str(payload.get("study_stage") or "Completed study").strip(),
                "research_area": str(payload.get("research_area") or "").strip(),
                "study_context": str(payload.get("context") or "").strip(),
                "citation_style": str(payload.get("citation_style") or "APA 7th").strip(),
                "school_format": str(payload.get("school_guidelines") or "").strip(),
                "background_structure": str(payload.get("background_structure") or "continuous_narrative").strip(),
                "purpose_statement_style": str(payload.get("purpose_statement_style") or "concise_general_objective").strip(),
                "target_pages": f"{page_min}-{page_max}",
                "target_citations_per_1000_words": f"{citation_min}-{citation_max}",
                "allow_missing_section_insertions": bool(payload.get("allow_missing_section_insertions", True)),
                "uploaded_content_scope": str(payload.get("uploaded_content_scope") or "selected_chapter"),
                "strengthening_scope": strengthening_scope,
                "selected_sections_to_strengthen": selected_titles,
                "standard_sections_to_add_if_missing": new_titles,
                "custom_sections_to_add": custom_new_sections,
            },
            "long_chapter_strengthening_strategy": _revision_long_chapter_strategy(level, chapter_type, page_min, page_max),
            "previous_chapters_for_alignment": _previous_chapters_revision_context(payload),
            "research_logic": {
                "objectives": str(payload.get("objectives") or "").strip(),
                "research_questions": str(payload.get("research_questions") or "").strip(),
                "hypotheses": str(payload.get("hypotheses") or "").strip(),
                "theory_or_framework": str(payload.get("theory_framework") or "").strip(),
                "variables_or_constructs": str(payload.get("variables_constructs") or "").strip(),
                "methodology_and_analysis": str(payload.get("methodology") or "").strip(),
                "confirmed_results_or_findings": str(payload.get("data_and_results") or "").strip(),
                "known_contribution": str(payload.get("contribution_claim") or "").strip(),
            },
            "revision_direction": {
                "revision_level": str(payload.get("revision_level") or "Comprehensive chapter strengthening").strip(),
                "additional_goals": str(payload.get("revision_goals") or "").strip(),
                "supervisor_comments": str(payload.get("supervisor_comments") or "").strip(),
                "focus": _revision_focus(payload),
            },
            "existing_chapter": chapter_text,
            "section_scope": {
                "mode": strengthening_scope,
                "sections_to_strengthen": selected_titles,
                "sections_to_add_if_missing": new_titles,
                "custom_sections_to_add": custom_new_sections,
                "rules": (
                    [
                        "Return only the selected sections and requested new sections, not the full chapter.",
                        "Use the rest of the selected chapter only as context. Do not rewrite, summarise or reproduce unselected sections.",
                        "Preserve the selected chapter's numbering style and place requested additions in their academically appropriate location.",
                    ]
                    if strengthening_scope == "selected_sections"
                    else [
                        "Return the complete strengthened selected chapter only. Never return another chapter from a complete thesis upload.",
                        "Strengthen the complete selected chapter while preserving confirmed content and chapter boundaries.",
                    ]
                ),
            },
            "chapter_specific_rules": _chapter_rules(chapter_type, str(payload.get("study_stage") or "")),
            "scholarly_source_records": source_records,
            "strict_rules": [
                "Preserve confirmed facts, sample details, dates, quotations, coefficients, p-values, confidence intervals, table values, themes and study findings unless the user explicitly supplied a correction.",
                "Do not invent citations, references, data, analyses, instruments, ethics approvals, permissions, institutional rules, tables, figures or results.",
                "Retain valid existing citations. Do not alter author names or publication years unless the supplied evidence confirms a correction.",
                "Use retrieved metadata only when the title or abstract directly supports the claim. Never turn metadata into evidence for a result not reported by the source.",
                "Increase citation density according to the stronger level- and chapter-specific target, but avoid citation padding and do not attach citations to unsupported claims.",
                "For Chapter One, support most substantive background and problem paragraphs with at least one directly relevant source; evidence-heavy paragraphs may synthesise two or more sources.",
                "For Chapter Two, make nearly every substantive paragraph evidence-supported and normally compare two to four relevant sources in thematic synthesis rather than repeating one citation.",
                "Run a claim-evidence audit. Support every substantive factual, historical, policy, contextual, theoretical and empirical claim with a directly relevant and accurate source from the existing citations or verified source bank.",
                "For Chapter One, follow target_citations_per_1000_words in the project profile. Distribute accurate citations across the background, problem, significance and other evidence-led sections, while keeping objectives, questions, purpose and purely organisational sentences citation-light.",
                "Do not leave comments or user instructions inside the scholarly narrative. Put every unresolved action on a separate line beginning [ACTION REQUIRED: ...] immediately after the affected paragraph or list. Do not collect actions at the end of the chapter.",
                "Complete every correction that can be responsibly made from the chapter, project profile, earlier chapters and verified source bank. Do not create an action merely to ask the user to approve generated wording.",
                "Reserve ACTION REQUIRED items for unique material facts or institutional decisions that cannot be inferred or retrieved, such as the exact study population, site, sample, instrument, ethics approval, study period or actual results. State each missing input once at its first relevant location.",
                "For an Introduction chapter, if background_structure is continuous_narrative, keep Background to the Study under one main heading without internal numbered subheadings. Write Purpose of the Study concisely: one sentence when purpose_statement_style is concise_general_objective, or one short paragraph only when the school explicitly requires it. Do not add explanatory commentary after the purpose.",
                "Present research objectives and research questions as clean standalone lists. Remove explanatory commentary, level-alignment notes and methodological justification after the list.",
                "Restart research-question numbering at 1, independently of objective numbering. Use one numbered question for each item.",
                "End the revised chapter with one clean APA 7 References section containing only cited sources, one complete entry per paragraph, deduplicated and alphabetised. Do not use bullets, numbering, annotations, source keys, relevance labels or a Source Use Audit.",
                "Do not add mediation, moderation, causality, longitudinal design, multilevel structure, robustness tests or measurement validation unless they are justified by the approved study and available data.",
                "Do not present a recommended analysis as completed. Use a concise bracketed action item such as [conduct and report the required diagnostic test] when essential evidence is missing.",
                "Preserve chapter numbering and school-specific headings where supplied. Add a missing expected heading only when the chapter type, school guidelines, previous chapters or supervisor comments clearly require it.",
                "The uploaded input may be a complete thesis, but existing_chapter has already been isolated to the selected chapter. Do not output any other chapter.",
                "Obey section_scope exactly. In selected_sections mode, return only the chosen existing sections and requested additions. Do not revise or reproduce unselected sections.",
                "For a requested standard or custom new section, do not duplicate an equivalent section that already exists. Strengthen the existing equivalent instead and explain this in the report.",
                "Place [confirm added section: Section Title] immediately before every genuinely new section so the DOCX marks the addition for user confirmation.",
                "When an important section is missing and allow_missing_section_insertions is true, insert the missing heading or section in the revised chapter but mark it for confirmation using a bracketed red-action marker, for example [confirm added section: Theoretical Framework].",
                "If an added section requires evidence that was not supplied, include only defensible bridging prose and a precise bracketed placeholder such as [insert verified source for this added subsection] or [confirm supervisor approval for this added subsection].",
                "Use previous_chapters_for_alignment to check consistency from Chapter Two onward. Do not copy earlier chapters, but align variables, objectives, theory, method, terminology and scope.",
                "Strengthen paragraphs through claim, evidence, synthesis, interpretation, qualification and linkage to the study. Avoid repetitive templates and author-by-author listing.",
                *scholarly_humanizer_prompt_rules(),
                "Do not expand the chapter merely to hit a page target. Add depth through evidence, critique, theory, methodological justification, interpretation and cross-section alignment.",
                "When long_chapter_strengthening_strategy is enabled, treat the chapter as staged section strengthening. Do not compress a doctoral literature review into a short summary; diagnose and deepen conceptual, theoretical, empirical, methodological, contextual, gap and framework coverage separately.",
                "Use Markdown headings and tables where useful. Keep equations intact.",
                "Mark only genuine student or supervisor action items in square brackets so the complete instruction appears red in the DOCX, placed directly after the affected text.",
            ],
            "report_requirements": [
                "State whether the uploaded text matches the selected chapter type.",
                "Summarise the important changes made by section.",
                "Assess chapter structure, logical flow and alignment with the selected academic level.",
                "Assess the problem, gap, objectives, questions, hypotheses, theory, methods, results and conclusions only where relevant to this chapter.",
                "Report the achieved word count, estimated pages and citation density, and compare them with the planning targets without claiming that page count alone proves quality.",
                "List unresolved evidence, analysis, citation and author-action requirements.",
                "Identify cross-chapter consistency checks required before submission, especially issues detected from previous_chapters_for_alignment.",
                "List any added or proposed sections that were inserted for confirmation, and explain why each was needed.",
                "Do not guarantee supervisor approval, examination success or thesis acceptance.",
            ],
            "output_format": [
                "Return plain Markdown using the exact markers below.",
                (
                    "Start with ===REVISED_CHAPTER=== followed only by the strengthened selected sections and requested additions."
                    if strengthening_scope == "selected_sections"
                    else "Start with ===REVISED_CHAPTER=== followed by the complete strengthened selected chapter."
                ),
                "Then add ===STRENGTHENING_REPORT=== followed by the Chapter Strengthening Report.",
                "When supervisor comments were supplied and include_supervisor_response_matrix is true, add ===SUPERVISOR_RESPONSE_MATRIX=== followed by a Markdown table with columns Supervisor comment, Revision made, Location and Remaining action.",
                "Do not use code fences.",
            ],
            "include_supervisor_response_matrix": bool(payload.get("include_supervisor_response_matrix", True)),
        }

        revision_instructions = (
            "You are ProjectReady AI's senior thesis and dissertation chapter editor. "
            "Strengthen the supplied chapter rigorously while preserving confirmed evidence, valid citations, "
            "approved research logic and the student's substantive voice. Apply chapter-aware academic standards, "
            "formal British English and protected natural scholarly flow. Never invent analysis, results or references. "
            "Use clear discipline-specific wording rather than mechanical synonym replacement. "
            "Separate completed revisions from remaining student or supervisor actions."
        )
        try:
            raw, model_used, attempt_errors = _call_responses_with_fallbacks(
                client,
                level=level,
                prompt=prompt,
                instructions=revision_instructions,
                max_output_tokens=max(4000, min(_env_int("PROJECTREADY_CHAPTER_REVISION_MAX_OUTPUT_TOKENS", 18000, minimum=4000), 60000)),
            )
            provider_errors.extend(attempt_errors)
            revised_chapter, strengthening_report, supervisor_matrix = _split_revision_package(raw)
            if not revised_chapter:
                revised_chapter = chapter_text
            if not strengthening_report:
                strengthening_report = _fallback_report(
                    payload, source_records, ["The revision model returned no separate strengthening report."]
                )
            if payload.get("supervisor_comments") and payload.get("include_supervisor_response_matrix", True) and not supervisor_matrix:
                supervisor_matrix = _fallback_matrix(str(payload.get("supervisor_comments") or ""))
            mode = "ai_revision"
        except Exception as exc:
            provider_errors.append({"provider": "openai", "error": f"Chapter revision failed: {str(exc)[:220]}"})
            revised_chapter = chapter_text
            strengthening_report = _fallback_report(payload, source_records, provider_errors)
            supervisor_matrix = _fallback_matrix(str(payload.get("supervisor_comments") or ""))
            mode = "metadata_fallback_after_ai_error"

    humanizer_mode = str(payload.get("humanizer_mode") or os.getenv("PROJECTREADY_HUMANIZER_MODE", "balanced") or "balanced").strip().lower()
    if mode == "ai_revision":
        revised_chapter, humanizer_report = humanize_scholarly_text(revised_chapter, mode=humanizer_mode)
        humanizer_model = os.getenv("OPENAI_HUMANIZER_MODEL", "gpt-5.6-terra").strip() or model_used
        revised_chapter = _humanize_strengthened_chapter_with_model(
            client,
            humanizer_model,
            revised_chapter,
            mode=humanizer_mode,
            level=level,
            chapter_type=chapter_type,
            payload=payload,
        )
        revised_chapter, final_humanizer_report = humanize_scholarly_text(revised_chapter, mode=humanizer_mode)
        humanizer_report = {**humanizer_report, "final_score": final_humanizer_report.get("score")}
    else:
        humanizer_report = {"mode": humanizer_mode, "applied": False, "reason": "No completed AI revision to refine."}

    revised_chapter = _finalise_chapter_text(revised_chapter)
    revised_chapter = _normalise_purpose_of_study(revised_chapter)
    revised_chapter = _normalise_objectives_and_questions(revised_chapter)
    revised_chapter = _clean_chapter_references(revised_chapter)
    revised_chapter = _ensure_markdown_heading_spacing(revised_chapter)
    revised_chapter = detach_action_items(_finalise_chapter_text(revised_chapter))
    strengthening_report = _finalise_chapter_text(strengthening_report)
    supervisor_matrix = _finalise_chapter_text(supervisor_matrix) if supervisor_matrix else ""
    metrics = _metrics(revised_chapter)

    return {
        "revised_chapter_text": revised_chapter,
        "strengthening_report": strengthening_report,
        "supervisor_response_matrix": supervisor_matrix,
        "mode": mode,
        "model_used": model_used if client and ai_enabled else "none",
        "selected_level": level,
        "selected_chapter_type": chapter_type,
        "strengthening_scope": strengthening_scope,
        "selected_section_titles": selected_titles,
        "new_section_titles": new_titles,
        "custom_new_sections": custom_new_sections,
        "scope_metadata": scope_metadata,
        "processed_original_chapter_text": chapter_text,
        "target_page_range": f"{page_min}-{page_max}",
        "target_citation_density": f"{citation_min}-{citation_max} per 1,000 words",
        **metrics,
        "source_records_used": source_records,
        "source_bank_count": len(source_records),
        "excluded_retracted_count": len(blocked),
        "excluded_retracted_titles": [str(item.get("title") or "Untitled") for item in blocked[:10]],
        "provider_errors": provider_errors,
        "humanizer_report": humanizer_report,
        "revision_colour_note": (
            "In the downloaded DOCX, wording added or changed by ProjectReady AI is shown in blue. "
            "Student or supervisor action items are shown in red. Exact unchanged wording remains black."
        ),
        "quality_filters": [
            "The selected chapter type, requested section scope and academic level control the strengthening rules.",
            "A complete thesis upload is isolated to the selected chapter before revision.",
            "Selected-sections mode does not rewrite or return unselected sections.",
            "Confirmed numerical and qualitative evidence is preserved.",
            "Missing analysis is marked as action required rather than presented as completed.",
            "Citation density is strengthened through relevant evidence, not reference padding.",
            "Retracted or withdrawn records are excluded where detectable.",
            "The report does not guarantee supervisor approval or examination success.",
            "If the AI revision model is unavailable, the protected route returns an error instead of consuming paid revision entitlement.",
        ],
    }


def chapter_planning_targets(
    level: str,
    chapter_type: str,
    *,
    strengthening_scope: str = "whole_chapter",
    selected_section_count: int = 0,
    custom_target_pages_enabled: bool = False,
    target_page_min: int | None = None,
    target_page_max: int | None = None,
) -> dict[str, Any]:
    target_payload = {
        "strengthening_scope": strengthening_scope,
        "selected_section_titles": [f"Section {index + 1}" for index in range(max(0, int(selected_section_count or 0)))],
        "custom_target_pages_enabled": custom_target_pages_enabled,
        "target_page_min": target_page_min,
        "target_page_max": target_page_max,
    }
    page_min, page_max = _resolved_page_target(target_payload, level, chapter_type)
    citation_min, citation_max = _citation_target(level, chapter_type)
    return {
        "academic_level": level,
        "chapter_type": chapter_type,
        "page_range": {"minimum": page_min, "maximum": page_max},
        "citation_density_per_1000_words": {"minimum": citation_min, "maximum": citation_max},
        "word_range_estimate": {"minimum": page_min * 330, "maximum": page_max * 380},
        "strengthening_scope": strengthening_scope,
        "custom_target_applied": bool(custom_target_pages_enabled),
        "note": "Page and citation ranges are planning targets for the selected strengthening scope. Quality, institutional requirements and evidence availability remain controlling considerations.",
    }



def _plain_compare_text(text: str) -> str:
    text = re.sub(r"^#{1,6}\s+", "", str(text or "").strip())
    text = re.sub(r"^[-*•]\s+", "", text)
    text = re.sub(r"^\d+[.)]\s+", "", text)
    text = _plain_inline_text(text)
    text = re.sub(r"\s+", " ", text)
    return text.strip().lower()


def _tokenise_for_diff(text: str) -> list[str]:
    return re.findall(r"\s+|[A-Za-z0-9]+(?:['’-][A-Za-z0-9]+)*|[^\w\s]", text, flags=re.UNICODE)


def _word_key(token: str) -> str:
    if token.isspace():
        return token
    return token.lower()


def _best_original_line(revised_line: str, original_lines: list[str], used: set[int]) -> tuple[str | None, float, int | None]:
    revised_norm = _plain_compare_text(revised_line)
    if not revised_norm:
        return None, 0.0, None
    best_line = None
    best_score = 0.0
    best_index = None
    for index, candidate in enumerate(original_lines):
        if index in used:
            continue
        candidate_norm = _plain_compare_text(candidate)
        if not candidate_norm:
            continue
        score = difflib.SequenceMatcher(None, candidate_norm, revised_norm, autojunk=False).ratio()
        if score > best_score:
            best_score = score
            best_line = candidate
            best_index = index
    return best_line, best_score, best_index


def _append_marked_run(
    paragraph,
    text: str,
    changed: bool,
    bold: bool = False,
    italic: bool = False,
    action_required: bool = False,
) -> None:
    if not text:
        return
    from docx.shared import RGBColor

    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if action_required:
        run.font.color.rgb = RGBColor(*_ACTION_RED)
    elif changed:
        run.font.color.rgb = RGBColor(*_REVISION_BLUE)


def _styled_diff_tokens(text: str) -> list[tuple[str, bool, bool, bool]]:
    tokens: list[tuple[str, bool, bool, bool]] = []
    for visible, bold, italic, action_required in _parse_inline_segments(text):
        for token in _tokenise_for_diff(visible):
            tokens.append((token, bold, italic, action_required))
    return tokens


def _add_diff_runs(paragraph, revised_text: str, original_text: str | None, changed_default: bool = True) -> None:
    """Render emphasis, keep revisions blue and override action-required text in red."""
    revised_tokens = _styled_diff_tokens(revised_text)
    if not original_text:
        for token, bold, italic, action_required in revised_tokens:
            _append_marked_run(
                paragraph, token, True if token else changed_default, bold, italic, action_required
            )
        return

    original_tokens = _styled_diff_tokens(original_text)
    matcher = difflib.SequenceMatcher(
        None,
        [_word_key(token[0]) for token in original_tokens],
        [_word_key(token[0]) for token in revised_tokens],
        autojunk=False,
    )
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        for offset, (token, bold, italic, action_required) in enumerate(revised_tokens[j1:j2]):
            changed = tag != "equal"
            if tag == "equal" and i1 + offset < i2:
                original_token = original_tokens[i1 + offset]
                # A formatting change is also a revision and should remain blue.
                if not token.isspace():
                    changed = (bold, italic) != (original_token[1], original_token[2])
            _append_marked_run(paragraph, token, changed, bold, italic, action_required)


def _add_black_inline_runs(paragraph, text: str, force_bold: bool = False) -> None:
    """Render Markdown emphasis in black, with action-required text in red."""
    for visible, bold, italic, action_required in _parse_inline_segments(text):
        _append_marked_run(
            paragraph,
            visible,
            False,
            bool(bold or force_bold),
            italic,
            action_required,
        )

def _revision_line_map(original_text: str, revised_lines: list[str]) -> dict[int, str | None]:
    original_lines = [line.rstrip() for line in _finalise_article_text(original_text).splitlines() if line.strip()]
    used: set[int] = set()
    mapping: dict[int, str | None] = {}
    exact_lookup: dict[str, list[int]] = {}
    for index, line in enumerate(original_lines):
        exact_lookup.setdefault(_plain_compare_text(line), []).append(index)

    for revised_index, line in enumerate(revised_lines):
        norm = _plain_compare_text(line)
        exact_candidates = exact_lookup.get(norm, [])
        exact_index = next((idx for idx in exact_candidates if idx not in used), None)
        if exact_index is not None:
            used.add(exact_index)
            mapping[revised_index] = original_lines[exact_index]
            continue
        candidate, score, candidate_index = _best_original_line(line, original_lines, used)
        if candidate is not None and score >= 0.36 and candidate_index is not None:
            used.add(candidate_index)
            mapping[revised_index] = candidate
        else:
            mapping[revised_index] = None
    return mapping


def _add_revision_table(doc, lines: list[str], original_text: str, colour_revisions: bool = True) -> None:
    rows: list[list[str]] = []
    for line in lines:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        if cells and not all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells):
            rows.append(cells)
    if not rows:
        return
    width = max(len(row) for row in rows)
    table = doc.add_table(rows=0, cols=width)
    table.style = "Table Grid"
    original_norm = _plain_compare_text(original_text)
    for row_index, cells in enumerate(rows):
        row = table.add_row().cells
        for column in range(width):
            value = cells[column] if column < len(cells) else ""
            paragraph = row[column].paragraphs[0]
            paragraph.clear()
            if colour_revisions:
                visible_value = _plain_compare_text(value)
                changed = bool(visible_value and visible_value not in original_norm)
                for visible, bold, italic, action_required in _parse_inline_segments(value):
                    _append_marked_run(
                        paragraph,
                        visible,
                        changed,
                        bool(bold or row_index == 0),
                        italic,
                        action_required,
                    )
            else:
                _add_black_inline_runs(paragraph, value, force_bold=row_index == 0)

def _write_markdown_document(doc, markdown_text: str, original_text: str | None = None, colour_revisions: bool = False, heading_offset: int = 0) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    lines = _finalise_article_text(markdown_text).splitlines()
    line_map = _revision_line_map(original_text or "", [line for line in lines]) if colour_revisions else {}
    table_buffer: list[str] = []
    code_buffer: list[str] = []
    in_code = False
    equation_buffer: list[str] = []
    in_equation = False

    def add_content(paragraph, content: str, index: int, original_override: str | None = None) -> None:
        original_line = original_override if original_override is not None else line_map.get(index)
        if colour_revisions:
            _add_diff_runs(paragraph, content, original_line)
        else:
            _add_black_inline_runs(paragraph, content)

    for index, raw_line in enumerate(lines):
        line = raw_line.rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                for code_line in code_buffer:
                    paragraph = doc.add_paragraph()
                    run = paragraph.add_run(code_line)
                    run.font.name = "Consolas"
                    run.font.size = Pt(9)
                    if colour_revisions:
                        run.font.color.rgb = RGBColor(*_REVISION_BLUE)
                code_buffer = []
                in_code = False
            else:
                if table_buffer:
                    _add_revision_table(doc, table_buffer, original_text or "", colour_revisions)
                    table_buffer = []
                in_code = True
            continue
        if in_code:
            code_buffer.append(line)
            continue

        if stripped == "$$":
            if in_equation:
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run(" ".join(equation_buffer))
                run.font.name = "Cambria Math"
                run.font.size = Pt(12)
                if colour_revisions and _plain_compare_text(run.text) not in _plain_compare_text(original_text or ""):
                    run.font.color.rgb = RGBColor(*_REVISION_BLUE)
                equation_buffer = []
                in_equation = False
            else:
                if table_buffer:
                    _add_revision_table(doc, table_buffer, original_text or "", colour_revisions)
                    table_buffer = []
                in_equation = True
            continue
        if in_equation:
            equation_buffer.append(stripped)
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            table_buffer.append(line)
            continue
        if table_buffer:
            _add_revision_table(doc, table_buffer, original_text or "", colour_revisions)
            table_buffer = []

        if not stripped:
            continue
        if line.startswith("# "):
            level = min(3, 0 + heading_offset)
            paragraph = doc.add_heading("", level=level)
            if level == 0:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_content(paragraph, line[2:].strip(), index)
        elif line.startswith("## "):
            paragraph = doc.add_heading("", level=min(3, 1 + heading_offset))
            add_content(paragraph, line[3:].strip(), index)
        elif line.startswith("### "):
            paragraph = doc.add_heading("", level=min(3, 2 + heading_offset))
            add_content(paragraph, line[4:].strip(), index)
        elif line.startswith("#### "):
            paragraph = doc.add_heading("", level=3)
            add_content(paragraph, line[5:].strip(), index)
        elif re.match(r"^[-*•]\s+", line):
            content = re.sub(r"^[-*•]\s+", "", line).strip()
            paragraph = doc.add_paragraph(style="List Bullet")
            add_content(paragraph, content, index)
        elif re.match(r"^\d+[.)]\s+", line):
            content = re.sub(r"^\d+[.)]\s+", "", line).strip()
            paragraph = doc.add_paragraph(style="List Number")
            add_content(paragraph, content, index)
        else:
            paragraph = doc.add_paragraph()
            add_content(paragraph, line, index)

    if table_buffer:
        _add_revision_table(doc, table_buffer, original_text or "", colour_revisions)
    if code_buffer:
        for code_line in code_buffer:
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(code_line)
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            if colour_revisions:
                run.font.color.rgb = RGBColor(*_REVISION_BLUE)
    if equation_buffer:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(" ".join(equation_buffer))
        run.font.name = "Cambria Math"
        run.font.size = Pt(12)
        if colour_revisions:
            run.font.color.rgb = RGBColor(*_REVISION_BLUE)


def export_revised_chapter_docx(
    original_chapter_text: str,
    revised_chapter_text: str,
    title: str = "Strengthened Thesis Chapter",
    strengthening_report: str = "",
    supervisor_response_matrix: str = "",
    include_strengthening_report: bool = True,
) -> tuple[io.BytesIO, str]:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    safe_title = re.sub(r"[^A-Za-z0-9_-]+", "_", (title or "strengthened_chapter")[:80]).strip("_") or "strengthened_chapter"
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.color.rgb = RGBColor(0, 0, 0)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run("Responsible-use notice: This is an editable AI-assisted working revision based on the user's existing chapter and research inputs. It is not a completed or submission-ready academic work. Wording added or changed by ProjectReady AI appears in blue, student or supervisor action items appear in red, and exact unchanged wording remains black. Verify every source, fact, method and finding before submission.")
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(89, 98, 115)

    _write_markdown_document(
        doc,
        revised_chapter_text,
        original_text=original_chapter_text,
        colour_revisions=True,
    )

    if include_strengthening_report and (strengthening_report.strip() or supervisor_response_matrix.strip()):
        doc.add_page_break()
        heading = doc.add_heading("Chapter Strengthening Report", level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if strengthening_report.strip():
            clean_report = re.sub(r"^#\s+Chapter Strengthening Report\s*", "", strengthening_report.strip(), count=1, flags=re.IGNORECASE)
            _write_markdown_document(doc, clean_report, colour_revisions=False, heading_offset=1)
        if supervisor_response_matrix.strip():
            doc.add_heading("Response to Supervisor Comments", level=1)
            _write_markdown_document(doc, supervisor_response_matrix, colour_revisions=False, heading_offset=1)

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream, f"{safe_title}_strengthened_working_revision.docx"
