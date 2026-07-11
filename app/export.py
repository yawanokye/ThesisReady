
from __future__ import annotations

import re
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def _safe_filename(value: str, fallback: str = "project") -> str:
    value = re.sub(r"[^A-Za-z0-9_-]+", "_", str(value or fallback)).strip("_")
    return (value or fallback)[:70]


def _set_cell_shading(cell, fill: str = "D9EAF7") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    _add_runs(p, text, bold_default=bold)


def _apply_document_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    for style_name in ["Normal", "Body Text"]:
        if style_name in styles:
            style = styles[style_name]
            style.font.name = "Times New Roman"
            style.font.size = Pt(12)
            try:
                style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            except Exception:
                pass

    for style_name in ["Heading 1", "Heading 2", "Heading 3", "Title"]:
        if style_name in styles:
            style = styles[style_name]
            style.font.name = "Times New Roman"
            try:
                style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            except Exception:
                pass


def _set_paragraph_spacing(paragraph) -> None:
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(6)


def _attention_placeholder_pattern() -> re.Pattern:
    return re.compile(
        r"(\[(?:(?:ACTION\s+REQUIRED(?:\s+\d+)?\s*:\s*)|(?:insert|verify|confirm|provide|supply|complete|replace|check|add|update|obtain|state|specify|include|conduct|resolve|revise)\b)[^\]]*\])",
        flags=re.IGNORECASE,
    )


def _clean_inline_markup(text: str) -> str:
    text = str(text or "")
    text = re.sub(r"<span\s+[^>]*color\s*:\s*#?(?:c00000|ff0000|red)[^>]*>(.*?)</span>", r"\1", text, flags=re.I | re.S)
    text = re.sub(r"</?span[^>]*>", "", text, flags=re.I)
    text = text.replace("[[ADD]]", "").replace("[[/ADD]]", "")
    return _clean_supplementary_table_noise(text)

def _clean_supplementary_table_noise(text: str) -> str:
    """Remove style-transition tokens that should not appear inside tables or item codes."""
    text = str(text or "")
    text = re.sub(r"^\s*(Indeed|Conversely|Importantly|Still|Yet),\s+", "", text, flags=re.I)
    text = re.sub(r"\b(Indeed|Conversely|Importantly|Still|Yet),\s+(?=(SQ|A|COM|INV|COL|OUT|OC)\d+\b)", "", text, flags=re.I)
    text = re.sub(r"\s*That matters\.\s*", " ", text, flags=re.I)
    return text.strip()



def _add_plain_or_attention_runs(paragraph, token: str, *, bold: bool = False, italic: bool = False) -> None:
    for part in _attention_placeholder_pattern().split(token):
        if part == "":
            continue
        run = paragraph.add_run(part)
        run.bold = bold
        run.italic = italic
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        if _attention_placeholder_pattern().fullmatch(part):
            run.font.color.rgb = RGBColor(192, 0, 0)


def _add_runs(paragraph, text: str, *, bold_default: bool = False, italic_default: bool = False) -> None:
    """Add markdown-ish runs. Only bracketed attention placeholders are red."""
    text = _clean_inline_markup(text)
    # basic **bold** and *italic* parsing while preserving placeholders
    tokens = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text)
    for token in tokens:
        if token == "":
            continue
        bold = bold_default
        italic = italic_default
        if token.startswith("**") and token.endswith("**"):
            token = token[2:-2]
            bold = True
        elif token.startswith("*") and token.endswith("*"):
            token = token[1:-1]
            italic = True
        _add_plain_or_attention_runs(paragraph, token, bold=bold, italic=italic)


def _is_markdown_table(block: str) -> bool:
    lines = [ln.strip() for ln in block.splitlines() if ln.strip()]
    return len(lines) >= 2 and all("|" in ln for ln in lines[:2]) and re.search(r"^\s*\|?\s*:?-{3,}:?", lines[1]) is not None


def _parse_table_row(line: str) -> list[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [cell.strip() for cell in line.split("|")]


def _set_repeat_table_header(row) -> None:
    try:
        tr_pr = row._tr.get_or_add_trPr()
        tbl_header = OxmlElement("w:tblHeader")
        tbl_header.set(qn("w:val"), "true")
        tr_pr.append(tbl_header)
    except Exception:
        pass


def _add_markdown_table(doc: Document, block: str) -> None:
    lines = [ln.strip() for ln in block.splitlines() if ln.strip()]
    rows = [_parse_table_row(lines[0])]
    for ln in lines[2:]:
        if re.search(r"^\s*\|?\s*:?-{3,}:?", ln):
            continue
        rows.append(_parse_table_row(ln))
    rows = [[_clean_supplementary_table_noise(cell) for cell in row] for row in rows]
    if not rows or not rows[0]:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.autofit = True
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    _set_repeat_table_header(table.rows[0])
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            cell = table.cell(r_idx, c_idx)
            value = row[c_idx] if c_idx < len(row) else ""
            _set_cell_text(cell, value, bold=(r_idx == 0))
            for p in cell.paragraphs:
                _set_paragraph_spacing(p)
                for run in p.runs:
                    run.font.size = Pt(10 if cols >= 5 else 11)
            if r_idx == 0:
                _set_cell_shading(cell)
    doc.add_paragraph("")


def _add_code_block(doc: Document, block: str) -> None:
    # Strip markdown fence and language label.
    content = re.sub(r"^```[A-Za-z0-9_-]*\s*", "", block.strip())
    content = re.sub(r"\s*```$", "", content.strip())
    for line in content.splitlines() or [""]:
        p = doc.add_paragraph()
        _set_paragraph_spacing(p)
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(10)


def _add_block(doc: Document, block: str) -> None:
    block = block.strip()
    if not block:
        return
    if _is_markdown_table(block):
        _add_markdown_table(doc, block)
        return
    if block.startswith("```"):
        _add_code_block(doc, block)
        return
    if re.match(r"^\s*[-*_]{3,}\s*$", block):
        doc.add_paragraph("")
        return
    # Heading block only if the whole first line is a heading.
    lines = block.splitlines()
    if len(lines) == 1 and lines[0].lstrip().startswith("#"):
        level = min(max(len(re.match(r"^#+", lines[0].lstrip()).group(0)), 1), 4)
        text = re.sub(r"^#+\s*", "", lines[0].strip())
        p = doc.add_heading(text, level=min(level, 3))
        _set_paragraph_spacing(p)
        return

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("#"):
            level = min(max(len(re.match(r"^#+", stripped).group(0)), 1), 4)
            txt = re.sub(r"^#+\s*", "", stripped)
            p = doc.add_heading(txt, level=min(level, 3))
            _set_paragraph_spacing(p)
        elif re.match(r"^\d+[.)]\s+", stripped):
            p = doc.add_paragraph(style="List Number")
            _set_paragraph_spacing(p)
            _add_runs(p, re.sub(r"^\d+[.)]\s+", "", stripped))
        elif re.match(r"^[-*+]\s+", stripped):
            p = doc.add_paragraph(style="List Bullet")
            _set_paragraph_spacing(p)
            _add_runs(p, re.sub(r"^[-*+]\s+", "", stripped))
        elif stripped.startswith("$$") and stripped.endswith("$$"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_paragraph_spacing(p)
            _add_runs(p, stripped.strip("$ "))
        else:
            p = doc.add_paragraph()
            _set_paragraph_spacing(p)
            _add_runs(p, stripped)


def _markdown_to_docx(doc: Document, markdown_text: str) -> None:
    # Preserve code fences as blocks, then split remaining markdown into paragraph blocks.
    text = (markdown_text or "").replace("\r\n", "\n")
    pattern = re.compile(r"```.*?```", flags=re.DOTALL)
    pos = 0
    for m in pattern.finditer(text):
        before = text[pos:m.start()]
        for block in re.split(r"\n\s*\n", before):
            _add_block(doc, block)
        _add_block(doc, m.group(0))
        pos = m.end()
    rest = text[pos:]
    for block in re.split(r"\n\s*\n", rest):
        _add_block(doc, block)


def export_chapter_docx(project: dict[str, Any], chapter_number: int, draft: str, out_dir: Path) -> Path:
    out_dir.mkdir(parents=True, exist_ok=True)
    project_title = str(project.get("title") or (project.get("profile") or {}).get("title") or "ProjectReady_AI_Draft")
    safe_title = _safe_filename(project_title)
    path = out_dir / f"{safe_title}_chapter_{chapter_number}_working_draft.docx"

    doc = Document()
    _apply_document_defaults(doc)
    if int(chapter_number or 0) == 7:
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    title_p = doc.add_heading(project_title, level=0)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if int(chapter_number or 0) == 7:
        doc.add_paragraph("Supplementary methods and analysis guide generated with ProjectReady AI.")
        guide_note = "Note: This is a working guide for instrument development, coding, validation, analysis planning and appendix organisation. Verify all scale sources, item adaptations, data rules and supervisor requirements before use."
    else:
        doc.add_paragraph(f"Chapter {chapter_number} AI-assisted working draft developed with ProjectReady AI.")
        guide_note = "Responsible-use notice: This is an editable AI-assisted working draft developed from information supplied by the user. It is not a completed or submission-ready thesis, dissertation or assignment. The user must verify all sources, evidence, statistics, data, findings and institutional requirements, revise the text using independent academic judgement, and obtain appropriate supervisor or institutional approval before submission."
    note = doc.add_paragraph()
    _add_runs(note, guide_note, italic_default=True)
    doc.add_paragraph("")

    _markdown_to_docx(doc, draft or "")
    doc.save(path)
    return path


def export_compliance_docx(project: dict[str, Any], chapter_number: int, check: dict[str, Any], out_dir: Path) -> Path:
    out_dir.mkdir(parents=True, exist_ok=True)
    safe_title = _safe_filename(project.get("title") or (project.get("profile") or {}).get("title") or "project")
    path = out_dir / f"{safe_title}_chapter_{chapter_number}_compliance.docx"

    doc = Document()
    _apply_document_defaults(doc)
    doc.add_heading("ProjectReady AI Academic Development and Compliance Report", level=0)
    doc.add_paragraph(f"Project: {project.get('title') or (project.get('profile') or {}).get('title', '')}")
    doc.add_paragraph(f"Chapter: {chapter_number}")
    doc.add_paragraph(f"Compliance score: {check.get('score_percent', 0)}%")

    items = check.get("items", []) or []
    table = doc.add_table(rows=1, cols=5)
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    headers = ["Section", "Requirement", "Status", "Evidence", "Suggested action"]
    for i, h in enumerate(headers):
        _set_cell_text(table.rows[0].cells[i], h, bold=True)
        _set_cell_shading(table.rows[0].cells[i])
    for item in items:
        row = table.add_row().cells
        row[0].text = str(item.get("section_title", ""))
        row[1].text = str(item.get("requirement", ""))
        row[2].text = str(item.get("status", ""))
        row[3].text = str(item.get("evidence", ""))
        row[4].text = str(item.get("suggested_action", ""))
    doc.save(path)
    return path


def export_instrument_docx(project: dict[str, Any], chapter_number: int, out_dir: Path) -> Path:
    out_dir.mkdir(parents=True, exist_ok=True)
    profile = project.get("profile") or {}
    safe_title = _safe_filename(project.get("title") or profile.get("title") or "project")
    path = out_dir / f"{safe_title}_draft_instrument.docx"
    doc = Document()
    _apply_document_defaults(doc)
    doc.add_heading("Draft Data Collection Instrument", level=0)
    doc.add_paragraph(f"Project: {project.get('title') or profile.get('title', '')}")
    doc.add_paragraph("This export is generated from the current project profile and should be reviewed against the approved methodology and supervisor comments.")

    constructs = profile.get("constructs") or profile.get("variables") or []
    if isinstance(constructs, str):
        constructs = [x.strip() for x in re.split(r"\n|;|,", constructs) if x.strip()]
    if not constructs:
        constructs = ["[insert construct or variable]", "[insert construct or variable]"]

    table = doc.add_table(rows=1, cols=5)
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    for i, h in enumerate(["Section", "Construct/Variable", "Draft item", "Scale/Response", "Source/Notes"]):
        _set_cell_text(table.rows[0].cells[i], h, bold=True)
        _set_cell_shading(table.rows[0].cells[i])
    for idx, construct in enumerate(constructs, 1):
        row = table.add_row().cells
        row[0].text = f"Section {idx}"
        row[1].text = str(construct)
        row[2].text = f"[insert validated item(s) for {construct}]"
        row[3].text = "[insert scale, e.g., 5-point Likert]"
        row[4].text = "[insert verified scale source or supervisor-approved item note]"
    doc.save(path)
    return path


def export_methods_supplement_docx(project: dict[str, Any], chapter_number: int, out_dir: Path) -> Path:
    out_dir.mkdir(parents=True, exist_ok=True)
    profile = project.get("profile") or {}
    safe_title = _safe_filename(project.get("title") or profile.get("title") or "project")
    path = out_dir / f"{safe_title}_supplementary_methods.docx"
    doc = Document()
    _apply_document_defaults(doc)
    doc.add_heading("Supplementary Methods and Analysis Guide", level=0)
    doc.add_paragraph(f"Project: {project.get('title') or profile.get('title', '')}")
    doc.add_paragraph("This working guide supports instrument design, source and scale traceability, variable coding, data-source planning, analysis decisions and appendix organisation. It does not replace the main methodology chapter.")

    doc.add_heading("Objective-to-Measurement Alignment", level=1)
    objectives = profile.get("objectives") or profile.get("specific_objectives") or []
    if isinstance(objectives, str):
        objectives = [x.strip() for x in re.split(r"\n|;", objectives) if x.strip()]
    if not objectives:
        objectives = ["[insert research objective]"]
    table = doc.add_table(rows=1, cols=5)
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    headers = ["Objective", "Variable/Construct", "Indicator", "Data source/item", "Appendix location"]
    for i, h in enumerate(headers):
        _set_cell_text(table.rows[0].cells[i], h, bold=True)
        _set_cell_shading(table.rows[0].cells[i])
    for obj in objectives:
        row = table.add_row().cells
        row[0].text = str(obj)
        row[1].text = "[insert variable/construct]"
        row[2].text = "[insert operational indicator]"
        row[3].text = "[insert questionnaire item/data source]"
        row[4].text = "[insert appendix reference]"
    doc.save(path)
    return path
