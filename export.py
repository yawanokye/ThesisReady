from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt


def export_chapter_docx(project: dict[str, Any], chapter_number: int, draft: str, out_dir: Path) -> Path:
    out_dir.mkdir(parents=True, exist_ok=True)
    safe_title = re.sub(r"[^A-Za-z0-9_-]+", "_", project.get("title", "project"))[:60]
    path = out_dir / f"{safe_title}_chapter_{chapter_number}.docx"

    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)

    doc.add_heading(project.get("title", "ProjectReady AI Draft"), level=0)
    doc.add_paragraph(f"Chapter {chapter_number} draft generated with ProjectReady AI.")
    doc.add_paragraph("Note: Verify all citations, evidence, data, page numbers, and supervisor requirements before submission.")

    for block in [b.strip() for b in draft.split("\n\n") if b.strip()]:
        if block.startswith("# "):
            doc.add_heading(block.replace("#", "").strip(), level=1)
        elif block.startswith("## "):
            doc.add_heading(block.replace("#", "").strip(), level=2)
        elif block.startswith("### "):
            doc.add_heading(block.replace("#", "").strip(), level=3)
        else:
            for line in block.splitlines():
                stripped = line.strip()
                if not stripped:
                    continue
                if stripped.startswith("- "):
                    doc.add_paragraph(stripped[2:], style="List Bullet")
                elif re.match(r"^\d+\.\s+", stripped):
                    doc.add_paragraph(re.sub(r"^\d+\.\s+", "", stripped), style="List Number")
                else:
                    doc.add_paragraph(stripped)

    doc.save(path)
    return path


def export_compliance_docx(project: dict[str, Any], chapter_number: int, check: dict[str, Any], out_dir: Path) -> Path:
    out_dir.mkdir(parents=True, exist_ok=True)
    safe_title = re.sub(r"[^A-Za-z0-9_-]+", "_", project.get("title", "project"))[:60]
    path = out_dir / f"{safe_title}_chapter_{chapter_number}_compliance.docx"

    doc = Document()
    doc.add_heading("ProjectReady AI Compliance Report", level=0)
    doc.add_paragraph(f"Project: {project.get('title', '')}")
    doc.add_paragraph(f"Chapter: {chapter_number}")
    doc.add_paragraph(f"Compliance score: {check.get('score_percent', 0)}%")

    table = doc.add_table(rows=1, cols=5)
    hdr = table.rows[0].cells
    hdr[0].text = "Section"
    hdr[1].text = "Requirement"
    hdr[2].text = "Status"
    hdr[3].text = "Evidence"
    hdr[4].text = "Suggested action"

    for item in check.get("items", []):
        row = table.add_row().cells
        row[0].text = item.get("section_title", "")
        row[1].text = item.get("requirement", "")
        row[2].text = item.get("status", "")
        row[3].text = item.get("evidence", "")
        row[4].text = item.get("suggested_action", "")

    doc.save(path)
    return path
