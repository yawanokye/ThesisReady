from io import BytesIO

from docx import Document

from app.action_items import detach_action_items
from app.export import _add_runs


def test_actions_are_placed_immediately_after_affected_paragraph():
    text = (
        "1.1 Introduction\n\n"
        "Some details require confirmation: [confirm school type], and [insert policy source]. "
        "The study examines assessment.\n\n"
        "The next paragraph remains unchanged."
    )
    out = detach_action_items(text)
    blocks = out.split("\n\n")

    paragraph_index = blocks.index("The study examines assessment.")
    assert blocks[paragraph_index + 1].startswith("[ACTION REQUIRED 1:")
    assert blocks[paragraph_index + 2].startswith("[ACTION REQUIRED 2:")
    assert blocks[paragraph_index + 3] == "The next paragraph remains unchanged."
    assert "USER ACTIONS REQUIRED" not in out


def test_action_only_paragraph_stays_at_original_location():
    text = "1.2 Background\n\nEvidence-based paragraph.\n\n[insert current curriculum source]\n\nNext paragraph."
    out = detach_action_items(text)
    blocks = out.split("\n\n")
    assert blocks == [
        "1.2 Background",
        "Evidence-based paragraph.",
        "[ACTION REQUIRED 1: Insert current curriculum source.]",
        "Next paragraph.",
    ]


def test_action_items_are_deduplicated_at_first_relevant_location():
    out = detach_action_items("Text [insert source].\n\nMore text [insert source].")
    assert out.count("[ACTION REQUIRED") == 1
    assert out.split("\n\n")[1].startswith("[ACTION REQUIRED 1:")


def test_complete_action_instruction_is_red_in_docx_run():
    doc = Document()
    paragraph = doc.add_paragraph()
    _add_runs(paragraph, "[ACTION REQUIRED 1: Insert the verified policy source.]")
    assert len(paragraph.runs) == 1
    run = paragraph.runs[0]
    assert run.text == "[ACTION REQUIRED 1: Insert the verified policy source.]"
    assert run.font.color.rgb is not None
    assert str(run.font.color.rgb) == "C00000"


def test_repeated_population_and_location_actions_are_collapsed():
    text = (
        "Purpose paragraph. [insert study population].\n\n"
        "Problem paragraph. [confirm target respondent group].\n\n"
        "Context paragraph. [insert study location].\n\n"
        "Scope paragraph. [confirm study area]."
    )
    out = detach_action_items(text)
    assert out.count("[ACTION REQUIRED") == 2


def test_generic_confirmation_of_generated_objective_is_not_retained():
    out = detach_action_items(
        "1.5 General Objective\n\nTo examine financial literacy.\n\n[confirm approved general objective]"
    )
    assert "ACTION REQUIRED" not in out
